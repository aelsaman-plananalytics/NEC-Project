"""
Report Generator for NEC Contract Analysis System.

Converts analysis JSON into professional markdown and PDF reports.
"""

import os
import re
import html
import tempfile
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, List, Optional, Union
from textwrap import wrap
from io import BytesIO

# Primavera stores total float in hours; convert to working days for report display.
WORKING_HOURS_PER_DAY = 8


def _escape(s: str) -> str:
    """Escape for HTML display."""
    return html.escape(str(s)) if s is not None else ""


def _presentation_text(s: Any) -> str:
    """Presentation-only wording replacements for report output (no logic change)."""
    if s is None:
        return ""
    t = str(s)
    # Full sentence (e.g. Section B "What we saw"): plain, direct wording
    t = t.replace(
        "Different representations noted (date format or value); see notes for detail.",
        "The contract and programme show a slight difference in date or value here. See the notes below for details."
    )
    t = t.replace("Different representations noted", "Minor scheduling interpretation differences identified.")
    return t


def _delivery_risk_profile(critical_pct: float, under_40_pct: float) -> str:
    """Derive delivery risk level from % critical and % under 40 days float. Presentation only."""
    if critical_pct < 10 and under_40_pct < 15:
        return "Low"
    if critical_pct < 20:
        return "Moderate"
    return "Elevated"


def _delivery_risk_explanation(risk_level: str) -> str:
    """Return narrative for delivery risk level. Presentation only."""
    if risk_level == "Low":
        return "Schedule flexibility is distributed across the programme with limited concentration of critical activities."
    elif risk_level == "Moderate":
        return "Some concentration of critical or low-float activities indicates moderate delivery sensitivity."
    elif risk_level == "Elevated":
        return "A high proportion of critical or low-float activities increases sensitivity to delay."
    return ""


def _derive_overall_confidence(json_data: Dict[str, Any]) -> str:
    """
    Derive overall confidence from validation_result when section does not provide it.
    Presentation-only: blockers_count (aligned == False), assumption_count (explicit_assumption == True).
    Does not modify json_data or write to DB.
    """
    alignment = json_data.get("alignment") or {}
    scope_coverage = alignment.get("scope_coverage") or {}
    obligations_report = _safe_list(scope_coverage.get("obligations_report"))
    blockers_count = sum(1 for o in obligations_report if o.get("aligned") is False)
    assumption_count = sum(
        1 for o in obligations_report
        if "explicit_assumption" in o and o.get("explicit_assumption") is True
    )
    if blockers_count == 0 and assumption_count == 0:
        return "High"
    if blockers_count == 0:
        return "Medium"
    return "Low"


def _confidence_display(s: Any) -> str:
    """Presentation only: show confidence column as High, Medium, or Low."""
    if s is None:
        return "—"
    t = str(s).strip().lower()
    if "high" in t or t == "high":
        return "High"
    if "low" in t or "judgement" in t or t == "low":
        return "Low"
    return "Medium"


def _safe_dict(val: Any, default: Optional[dict] = None) -> dict:
    """Return val if it is a dict, else default (or empty dict)."""
    return val if isinstance(val, dict) else (default or {})


def _safe_list(val: Any) -> list:
    """Return val if it is a list, else empty list."""
    return val if isinstance(val, list) else []


def _safe_get(obj: Any, key: str, default: Any = None) -> Any:
    """Safe .get(): return default if obj is not a dict."""
    if isinstance(obj, dict):
        return obj.get(key, default)
    return default


def _safe_row(row: Any) -> dict:
    """Return row if dict, else a minimal dict from str(row) to avoid .get() on str."""
    if isinstance(row, dict):
        return row
    s = str(row) if row is not None else ""
    return {"requirement": s, "observation": s, "action": s, "message": s, "contract_scope": s,
            "programme_activities": [], "representation_status": "", "evidence_role": "", "notes": s, "reason": s, "constraint": s,
            "programme_evidence": s, "handling": "", "item": s, "contract": s, "programme": s,
            "finding_or_check": s, "source_clause": s, "source_type": s, "validation_basis": s,
            "contract_activity": s, "activity_or_gate": s, "when_required": s, "shown_in_programme": s,
            "evidence": s, "status": s, "outcome": s, "explanation": s}


def _cell_para(styles, text: str, font_size: int = 8, max_chars: Optional[int] = None) -> "Paragraph":
    """Wrap text in a ReportLab Paragraph so it wraps in table cells and does not overlap. Escapes XML entities."""
    if text is None:
        text = ""
    text = str(text).strip()
    if max_chars and len(text) > max_chars:
        text = text[: max_chars - 3] + "..."
    # Escape for ReportLab (Paragraph uses XML-like markup)
    text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    if not text:
        text = "—"
    try:
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.platypus import Paragraph
        cell_style = ParagraphStyle(
            "TableCell",
            parent=styles["Normal"],
            fontSize=font_size,
            leading=font_size + 2,
            spaceBefore=2,
            spaceAfter=2,
            wordWrap="CJK",
        )
        return Paragraph(text, cell_style)
    except Exception:
        return Paragraph(text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"), styles["Normal"])

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak, Image
    from reportlab.lib.utils import ImageReader
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False
    A4 = None
    SimpleDocTemplate = None
    Table = None
    TableStyle = None
    Paragraph = None
    Spacer = None
    PageBreak = None
    Image = None
    ImageReader = None
    getSampleStyleSheet = None
    ParagraphStyle = None
    inch = None
    colors = None

try:
    from jinja2 import Environment, FileSystemLoader
    JINJA2_AVAILABLE = True
except ImportError:
    JINJA2_AVAILABLE = False
    Environment = None
    FileSystemLoader = None

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    Document = None
    Inches = None
    Pt = None
    RGBColor = None
    WD_ALIGN_PARAGRAPH = None


def _extract_float_stats_from_validation(json_data: Dict[str, Any]) -> Optional[Dict[str, Any]]:
    """
    Extract float from validation_result (programme_summary.activities).
    Source data is in hours (e.g. Primavera); convert to working days (÷8) for all statistics and display.
    Ignore None, exclude negative. Presentation only.
    """
    ps = json_data.get("programme_summary") or {}
    activities = ps.get("activities") or []
    raw_floats: List[float] = []
    for act in activities:
        if not isinstance(act, dict):
            continue
        val = act.get("float") if act.get("float") is not None else act.get("total_float")
        if val is None:
            continue
        try:
            raw_float = float(val)
        except (TypeError, ValueError):
            continue
        if raw_float < 0:
            continue
        raw_floats.append(raw_float)

    if not raw_floats:
        return None

    # Float is always stored in hours; convert to days (8-hour workday) for report
    float_days_list = [round(r / WORKING_HOURS_PER_DAY, 1) for r in raw_floats]

    total_activities = len(float_days_list)
    critical_count = sum(1 for f in float_days_list if f == 0)
    under_40_count = sum(1 for f in float_days_list if f < 40)
    over_40_count = sum(1 for f in float_days_list if f >= 40)
    max_float = round(max(float_days_list), 1)
    sorted_f = sorted(float_days_list)
    mid = total_activities // 2
    median_float = round(
        (sorted_f[mid] if total_activities % 2 else (sorted_f[mid - 1] + sorted_f[mid]) / 2.0),
        1,
    )
    critical_pct = round(100.0 * critical_count / total_activities, 1) if total_activities else 0
    under_40_pct = round(100.0 * under_40_count / total_activities, 1) if total_activities else 0
    over_40_pct = round(100.0 * over_40_count / total_activities, 1) if total_activities else 0

    # 20-day interval bins: 0–20, 20–40, ..., 340–360, 360+
    bucket_labels = [f"{i}-{i+20}" for i in range(0, 360, 20)] + ["360+"]
    n_bins = len(bucket_labels)
    histogram = [0] * n_bins
    for f_val in float_days_list:
        if f_val >= 360:
            histogram[-1] += 1
        else:
            idx = int(f_val // 20)
            if idx < n_bins - 1:
                histogram[idx] += 1
            else:
                histogram[-1] += 1

    return {
        "total_activities": total_activities,
        "critical_count": critical_count,
        "under_40_count": under_40_count,
        "over_40_count": over_40_count,
        "max_float": max_float,
        "median_float": median_float,
        "critical_pct": critical_pct,
        "under_40_pct": under_40_pct,
        "over_40_pct": over_40_pct,
        "histogram_buckets": bucket_labels,
        "histogram_counts": histogram,
    }


def _render_float_histogram_png(float_stats: Dict[str, Any]) -> Optional[BytesIO]:
    """
    Generate Float Distribution Profile histogram. 20-day bin intervals; X = Float (days), Y = Number of activities.
    Agg backend, PNG buffer; monochrome/subtle blue, clean gridlines. Presentation only.
    """
    if not float_stats or not isinstance(float_stats, dict):
        return None
    buckets = float_stats.get("histogram_buckets") or []
    counts = float_stats.get("histogram_counts") or []
    if not buckets or not counts or len(buckets) != len(counts):
        return None
    try:
        import matplotlib  # type: ignore[import-untyped]
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt  # type: ignore[import-untyped]
    except ImportError as e:
        import logging
        logging.getLogger(__name__).warning("[Schedule Float] matplotlib not available: %s", e)
        return None
    fig, ax = plt.subplots(figsize=(6, 3.5))
    fig.patch.set_facecolor("white")
    ax.set_facecolor("white")
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.spines["left"].set_color("black")
    ax.spines["bottom"].set_color("black")
    ax.tick_params(colors="black")
    ax.set_title("Float Distribution Profile", fontsize=11, color="black")
    ax.yaxis.grid(True, color="#cccccc", linestyle="-", linewidth=0.5)
    ax.set_axisbelow(True)
    x_pos = list(range(len(buckets)))
    ax.bar(x_pos, counts, color="#5a6c7d", edgecolor="black", linewidth=0.5)
    ax.set_xticks(x_pos)
    ax.set_xticklabels(buckets, fontsize=8, rotation=45, ha="right")
    ax.set_xlabel("Float (days)", fontsize=10, color="black")
    ax.set_ylabel("Number of Activities", fontsize=10, color="black")
    ax.set_ylim(bottom=0)
    plt.tight_layout()
    buf = BytesIO()
    fig.savefig(buf, format="png", dpi=120, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


class ReportGenerator:
    """
    Generates professional NEC contract analysis reports in Markdown and PDF formats.
    """
    
    def __init__(self):
        """Initialize the report generator."""
        from app.runtime_paths import RUNTIME_DIR
        self.output_dir = RUNTIME_DIR
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
        # Initialize Jinja2 environment for HTML template rendering
        if JINJA2_AVAILABLE:
            template_dir = Path(__file__).parent / "templates"
            self.jinja_env = Environment(
                loader=FileSystemLoader(str(template_dir)),
                autoescape=True
            )
        else:
            self.jinja_env = None
    
    def generate_markdown(self, analysis_json: dict) -> str:
        """
        Convert the analysis JSON into a full, structured, professional report
        in markdown format.
        
        Must include:
        - Project summary
        - Scope summary
        - Programme-critical clauses
        - Missing/blank fields
        - Contract completeness classification
        - Professional recommendations
        - Tables when needed
        
        Args:
            analysis_json: The JSON output from /api/analyze_contract
            
        Returns:
            str: Complete markdown report
        """
        md_lines = []
        
        # Extract data from JSON
        project_name = analysis_json.get("project", "Unknown Project")
        metadata = analysis_json.get("metadata", {})
        analysis_timestamp = metadata.get("analysis_timestamp", datetime.now().isoformat())
        extracted_clauses = analysis_json.get("extracted_clauses", {})
        contract_completeness = analysis_json.get("contract_completeness", {})
        scope_items = analysis_json.get("scope_items", [])
        constraints = analysis_json.get("constraints", [])
        milestones = analysis_json.get("milestones", [])
        
        # Format timestamp for display
        try:
            dt = datetime.fromisoformat(analysis_timestamp.replace('Z', '+00:00'))
            formatted_date = dt.strftime("%B %d, %Y at %H:%M:%S")
        except:
            formatted_date = analysis_timestamp
        
        # ============================================================
        # 1. TITLE PAGE
        # ============================================================
        md_lines.append("# NEC Contract Analysis Report\n")
        md_lines.append(f"**Project:** {project_name}  \n")
        md_lines.append(f"**Date:** {formatted_date}\n")
        md_lines.append(f"**Document:** {metadata.get('filename', 'N/A')}\n")
        md_lines.append("\n---\n\n")
        
        # ============================================================
        # 2. EXECUTIVE SUMMARY
        # ============================================================
        md_lines.append("## Executive Summary\n\n")
        
        # Contract type
        doc_type = contract_completeness.get("document_type", "unknown")
        is_template = contract_completeness.get("is_template", False)
        
        if doc_type == "template":
            md_lines.append("This contract document has been classified as a **template** or **incomplete contract**. ")
            md_lines.append("Many programme-critical clauses are missing or blank.\n\n")
        elif doc_type == "partial":
            md_lines.append("This contract document has been classified as **partially complete**. ")
            md_lines.append("Some programme-critical clauses are filled, but others remain blank or missing.\n\n")
        elif doc_type == "complete":
            md_lines.append("This contract document has been classified as **complete**. ")
            md_lines.append("Most programme-critical clauses have been filled with actual values.\n\n")
        else:
            md_lines.append("This contract document has been analyzed for programme-critical clauses.\n\n")
        
        # Statistics
        filled_pct = contract_completeness.get("filled_percentage", 0.0)
        blank_pct = contract_completeness.get("blank_percentage", 0.0)
        missing_pct = 100.0 - filled_pct - blank_pct
        
        total_mandatory = contract_completeness.get("total_mandatory", 0)
        mandatory_filled = contract_completeness.get("mandatory_filled", 0)
        mandatory_blank = contract_completeness.get("mandatory_blank", 0)
        mandatory_missing = contract_completeness.get("mandatory_missing", 0)
        
        md_lines.append("### Contract Completeness Statistics\n\n")
        md_lines.append(f"- **Filled Clauses:** {mandatory_filled} / {total_mandatory} ({filled_pct:.1f}%)\n")
        md_lines.append(f"- **Blank Clauses:** {mandatory_blank} / {total_mandatory} ({blank_pct:.1f}%)\n")
        md_lines.append(f"- **Missing Clauses:** {mandatory_missing} / {total_mandatory} ({missing_pct:.1f}%)\n\n")
        
        # Scope summary
        md_lines.append("### Scope Summary\n\n")
        md_lines.append(f"- **Total Scope Items:** {len(scope_items)}\n")
        md_lines.append(f"- **Constraints Identified:** {len(constraints)}\n")
        md_lines.append(f"- **Milestones Identified:** {len(milestones)}\n\n")
        
        md_lines.append("---\n\n")
        
        # ============================================================
        # 3. PROGRAMME-CRITICAL CLAUSE TABLE
        # ============================================================
        md_lines.append("## Programme-Critical Clauses\n\n")
        md_lines.append("The following table shows the status of all programme-critical NEC clauses:\n\n")
        
        # Create table header
        md_lines.append("| Clause | Title | Status | Value | Page |\n")
        md_lines.append("|--------|-------|--------|-------|------|\n")
        
        # Sort clauses by number
        clause_numbers = sorted(extracted_clauses.keys(), key=lambda x: (
            int(x.split('.')[0]), 
            int(x.split('.')[1]) if '.' in x and x.split('.')[1].isdigit() else 0
        ))
        
        for clause_num in clause_numbers:
            clause_data = extracted_clauses[clause_num]
            title = clause_data.get("title", "N/A")
            status = clause_data.get("status", "unknown")
            value = clause_data.get("value", "")
            page = clause_data.get("page")
            
            # Format status with emoji
            status_display = {
                "filled": "✅ Filled",
                "blank": "⚪ Blank",
                "missing": "❌ Missing"
            }.get(status, status.title())
            
            # Truncate long values
            if value:
                if len(value) > 80:
                    value_display = value[:77] + "..."
                else:
                    value_display = value
                # Escape pipe characters in markdown
                value_display = value_display.replace("|", "\\|")
            else:
                value_display = "*Not provided*"
            
            # Format page number
            page_display = str(page) if page is not None else "N/A"
            
            md_lines.append(f"| {clause_num} | {title} | {status_display} | {value_display} | {page_display} |\n")
        
        md_lines.append("\n---\n\n")
        
        # ============================================================
        # 4. CONTRACT COMPLETENESS ASSESSMENT
        # ============================================================
        md_lines.append("## Contract Completeness Assessment\n\n")
        
        # Classification
        md_lines.append(f"**Classification:** {doc_type.upper()}\n\n")
        md_lines.append(f"**Completeness Score:** {filled_pct:.1f}%\n\n")
        
        # Missing fields list
        missing_fields = contract_completeness.get("missing_fields", [])
        if missing_fields:
            md_lines.append("### Missing Clauses\n\n")
            md_lines.append("The following programme-critical clauses are missing from the contract:\n\n")
            for clause_num in sorted(missing_fields):
                clause_data = extracted_clauses.get(clause_num, {})
                title = clause_data.get("title", "Unknown")
                md_lines.append(f"- **{clause_num}** - {title}\n")
            md_lines.append("\n")
        
        # Blank fields list
        blank_fields = contract_completeness.get("blank_fields", [])
        if blank_fields:
            md_lines.append("### Blank Clauses\n\n")
            md_lines.append("The following programme-critical clauses are present but blank:\n\n")
            for clause_num in sorted(blank_fields):
                clause_data = extracted_clauses.get(clause_num, {})
                title = clause_data.get("title", "Unknown")
                md_lines.append(f"- **{clause_num}** - {title}\n")
            md_lines.append("\n")
        
        # Filled fields list
        filled_fields = contract_completeness.get("filled_fields", [])
        if filled_fields:
            md_lines.append("### Filled Clauses\n\n")
            md_lines.append("The following programme-critical clauses have been filled:\n\n")
            for clause_num in sorted(filled_fields):
                clause_data = extracted_clauses.get(clause_num, {})
                title = clause_data.get("title", "Unknown")
                value = clause_data.get("value", "")
                if len(value) > 100:
                    value = value[:97] + "..."
                md_lines.append(f"- **{clause_num}** - {title}: `{value}`\n")
            md_lines.append("\n")
        
        md_lines.append("---\n\n")
        
        # ============================================================
        # 5. RECOMMENDATIONS
        # ============================================================
        md_lines.append("## Recommendations\n\n")
        
        recommendations = self._generate_recommendations(
            contract_completeness,
            extracted_clauses,
            metadata
        )
        
        for i, rec in enumerate(recommendations, 1):
            md_lines.append(f"### {i}. {rec['title']}\n\n")
            md_lines.append(f"{rec['description']}\n\n")
            if rec.get('actions'):
                md_lines.append("**Recommended Actions:**\n\n")
                for action in rec['actions']:
                    md_lines.append(f"- {action}\n")
                md_lines.append("\n")
        
        md_lines.append("---\n\n")
        
        # ============================================================
        # 6. ADDITIONAL INFORMATION
        # ============================================================
        md_lines.append("## Additional Information\n\n")
        
        # Contract dates
        contract_dates = analysis_json.get("contract_dates", {})
        if contract_dates:
            md_lines.append("### Contract Dates\n\n")
            if contract_dates.get("starting_date"):
                md_lines.append(f"- **Starting Date:** {contract_dates['starting_date']}\n")
            if contract_dates.get("completion_date"):
                md_lines.append(f"- **Completion Date:** {contract_dates['completion_date']}\n")
            if contract_dates.get("possession_dates"):
                md_lines.append(f"- **Possession Dates:** {', '.join(contract_dates['possession_dates'])}\n")
            md_lines.append("\n")
        
        # Programme requirements
        programme_reqs = analysis_json.get("programme_requirements", {})
        if programme_reqs:
            md_lines.append("### Programme Requirements\n\n")
            if programme_reqs.get("submit_first_programme_within"):
                md_lines.append(f"- **First Programme Submission:** {programme_reqs['submit_first_programme_within']}\n")
            if programme_reqs.get("revised_programme_interval"):
                md_lines.append(f"- **Revised Programme Interval:** {programme_reqs['revised_programme_interval']}\n")
            md_lines.append("\n")
        
        # Delay damages
        delay_damages = analysis_json.get("delay_damages", "")
        if delay_damages:
            md_lines.append(f"### Delay Damages\n\n{delay_damages}\n\n")
        
        # Defects
        defects = analysis_json.get("defects", {})
        if defects:
            md_lines.append("### Defects Information\n\n")
            if defects.get("defects_date"):
                md_lines.append(f"- **Defects Date:** {defects['defects_date']}\n")
            if defects.get("defect_correction_period"):
                md_lines.append(f"- **Defect Correction Period:** {defects['defect_correction_period']}\n")
            if defects.get("landscaping_maintenance_period"):
                md_lines.append(f"- **Landscaping Maintenance Period:** {defects['landscaping_maintenance_period']}\n")
            md_lines.append("\n")
        
        # Payment terms
        payment_terms = analysis_json.get("payment_terms", {})
        if payment_terms:
            md_lines.append("### Payment Terms\n\n")
            if payment_terms.get("assessment_interval"):
                md_lines.append(f"- **Assessment Interval:** {payment_terms['assessment_interval']}\n")
            if payment_terms.get("payment_period"):
                md_lines.append(f"- **Payment Period:** {payment_terms['payment_period']}\n")
            if payment_terms.get("retention_percentage"):
                md_lines.append(f"- **Retention Percentage:** {payment_terms['retention_percentage']}\n")
            if payment_terms.get("bond_amount"):
                md_lines.append(f"- **Bond Amount:** {payment_terms['bond_amount']}\n")
            md_lines.append("\n")
        
        # Weather data
        weather_data = analysis_json.get("weather_data", {})
        if weather_data:
            md_lines.append("### Weather Data\n\n")
            if weather_data.get("recording_location"):
                md_lines.append(f"- **Recording Location:** {weather_data['recording_location']}\n")
            if weather_data.get("measurement_data"):
                md_lines.append(f"- **Measurement Data:** {weather_data['measurement_data']}\n")
            if weather_data.get("historical_records_source"):
                md_lines.append(f"- **Historical Records Source:** {weather_data['historical_records_source']}\n")
            md_lines.append("\n")
        
        # Metadata
        md_lines.append("### Report Metadata\n\n")
        md_lines.append(f"- **File Name:** {metadata.get('filename', 'N/A')}\n")
        md_lines.append(f"- **File Size:** {metadata.get('file_size_bytes', 0) / 1024 / 1024:.2f} MB\n")
        md_lines.append(f"- **Total Pages:** {metadata.get('total_pages', 0)}\n")
        md_lines.append(f"- **TOC Detected:** {'Yes' if metadata.get('toc_detected') else 'No'}\n")
        md_lines.append(f"- **Extraction Confidence:** {metadata.get('extraction_confidence', 0.0) * 100:.1f}%\n")
        
        if metadata.get('missing_sections'):
            md_lines.append(f"- **Missing Sections:** {', '.join(metadata['missing_sections'])}\n")
        
        md_lines.append("\n---\n\n")
        md_lines.append(f"*Report generated on {formatted_date}*\n")
        
        return "\n".join(md_lines)
    
    def _generate_recommendations(
        self,
        contract_completeness: Dict[str, Any],
        extracted_clauses: Dict[str, Any],
        metadata: Dict[str, Any]
    ) -> List[Dict[str, Any]]:
        """
        Generate professional recommendations based on contract analysis.
        
        Args:
            contract_completeness: Contract completeness data
            extracted_clauses: Extracted clause data
            metadata: Analysis metadata
            
        Returns:
            List of recommendation dictionaries with 'title', 'description', and 'actions'
        """
        recommendations = []
        
        doc_type = contract_completeness.get("document_type", "unknown")
        missing_fields = contract_completeness.get("missing_fields", [])
        blank_fields = contract_completeness.get("blank_fields", [])
        filled_pct = contract_completeness.get("filled_percentage", 0.0)
        
        # Recommendation 1: Missing clauses
        if missing_fields:
            recommendations.append({
                "title": "Address Missing Programme-Critical Clauses",
                "description": (
                    f"The contract is missing {len(missing_fields)} programme-critical clause(s). "
                    "These clauses are essential for project planning and risk management. "
                    "Without these values, it will be difficult to create an accurate programme "
                    "and assess project risks."
                ),
                "actions": [
                    "Review the contract document to locate the missing clauses",
                    "Consult with the Employer or Project Manager to obtain the missing values",
                    "Update the contract document with the missing clause values",
                    "Re-run the analysis to verify all clauses are now present"
                ]
            })
        
        # Recommendation 2: Blank clauses
        if blank_fields:
            recommendations.append({
                "title": "Complete Blank Clause Fields",
                "description": (
                    f"The contract contains {len(blank_fields)} blank programme-critical clause(s). "
                    "While these clauses are present in the document, they have not been filled with "
                    "actual values. This may indicate an incomplete contract or a template that needs "
                    "to be populated."
                ),
                "actions": [
                    "Identify the blank clause fields in the contract document",
                    "Determine the appropriate values for each blank clause",
                    "Fill in the blank fields with the correct values",
                    "Ensure all values are consistent with project requirements"
                ]
            })
        
        # Recommendation 3: Low completeness
        if filled_pct < 50.0:
            recommendations.append({
                "title": "Improve Contract Completeness",
                "description": (
                    f"The contract has a completeness score of {filled_pct:.1f}%, which is below the "
                    "recommended threshold. A complete contract should have at least 80% of "
                    "programme-critical clauses filled. This low score may indicate significant "
                    "risks for project planning and execution."
                ),
                "actions": [
                    "Prioritize filling the most critical clauses (dates, deadlines, payment terms)",
                    "Establish a contract completion plan with clear deadlines",
                    "Assign responsibility for completing each missing or blank clause",
                    "Regularly review and update contract completeness status"
                ]
            })
        
        # Recommendation 4: Template detection
        if doc_type == "template" or contract_completeness.get("is_template", False):
            recommendations.append({
                "title": "Convert Template to Active Contract",
                "description": (
                    "This contract has been identified as a template or incomplete document. "
                    "Templates should not be used for active projects without proper completion. "
                    "Using an incomplete contract can lead to disputes, delays, and legal issues."
                ),
                "actions": [
                    "Verify that this is the correct contract document for the project",
                    "Complete all programme-critical clauses before project commencement",
                    "Obtain formal approval from all parties before using this contract",
                    "Consider using a contract completion checklist to ensure nothing is missed"
                ]
            })
        
        # Recommendation 5: Legal/engineering implications
        critical_missing = [
            "3.1",  # Starting Date
            "3.3",  # Completion Date
            "3.7",  # Delay Damages
            "5.5",  # Retention Percentage
        ]
        
        missing_critical = [c for c in critical_missing if c in missing_fields or c in blank_fields]
        if missing_critical:
            recommendations.append({
                "title": "Address Critical Legal and Engineering Implications",
                "description": (
                    f"The contract is missing {len(missing_critical)} critical clause(s) that have "
                    "significant legal and engineering implications. These clauses define key "
                    "project parameters, financial terms, and risk allocation."
                ),
                "actions": [
                    "Immediately address missing Starting Date (3.1) and Completion Date (3.3)",
                    "Clarify Delay Damages (3.7) to understand financial risk exposure",
                    "Confirm Retention Percentage (5.5) to understand cash flow implications",
                    "Consult with legal and project management teams before proceeding"
                ]
            })
        
        # Recommendation 6: Next steps
        if not recommendations:
            recommendations.append({
                "title": "Contract Review Complete",
                "description": (
                    "The contract appears to be in good condition with most programme-critical "
                    "clauses filled. Continue with standard project planning and execution processes."
                ),
                "actions": [
                    "Proceed with programme development using the extracted clause values",
                    "Monitor contract compliance throughout the project lifecycle",
                    "Update the contract analysis if any changes are made to the contract"
                ]
            })
        else:
            recommendations.append({
                "title": "Next Steps",
                "description": (
                    "After addressing the recommendations above, proceed with the following steps:"
                ),
                "actions": [
                    "Re-run the contract analysis to verify improvements",
                    "Use the extracted clause values to develop the project programme",
                    "Set up monitoring and compliance checks for contract requirements",
                    "Schedule regular contract review meetings"
                ]
            })
        
        return recommendations
    
    def generate_pdf(
        self,
        json_data: Dict[str, Any],
        output_path: Optional[str] = None,
        report_options: Optional[Dict[str, Any]] = None,
    ) -> Union[bytes, str]:
        """
        Generates a professional PDF report from JSON data.
        If input is validation JSON (alignment + contract_summary), produces Programme Validation Report (Sections A–G).
        Otherwise uses contract analysis template.
        Falls back to DOCX if PDF generation fails.
        report_options: optional dict with confidentiality_mode (redact activity names), organisation_logo_url, user_name.
        """
        opts = report_options or json_data.pop("_report_options", None) or {}
        is_validation = "alignment" in json_data and "contract_summary" in json_data
        if is_validation and REPORTLAB_AVAILABLE:
            try:
                return self._generate_validation_pdf(json_data, output_path, report_options=opts)
            except Exception as e:
                print(f"[ReportGenerator] Validation PDF failed: {e}, attempting DOCX fallback...")
                if DOCX_AVAILABLE:
                    return self._generate_validation_docx(json_data, output_path)
                raise
        try:
            if not REPORTLAB_AVAILABLE:
                raise ImportError("ReportLab not available")
            
            # Prepare data for template rendering (contract analysis)
            template_data = self._prepare_template_data(json_data)
            
            # Generate PDF using ReportLab
            buffer = BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=A4, 
                                    rightMargin=72, leftMargin=72,
                                    topMargin=72, bottomMargin=72)
            
            # Build PDF content
            story = []
            styles = getSampleStyleSheet()
            # Table cell style so text wraps and does not overlap (leading/space avoid row height underestimate)
            table_cell_style = ParagraphStyle(
                "TableCell",
                parent=styles["Normal"],
                fontSize=7,
                leading=11,
                spaceBefore=4,
                spaceAfter=4,
                wordWrap="CJK",
            )
            def _p(s: str, max_len: Optional[int] = None) -> "Paragraph":
                if s is None:
                    s = "—"
                s = str(s).strip()
                if max_len and len(s) > max_len:
                    s = s[: max_len - 3] + "..."
                s = s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;") or "—"
                return Paragraph(s, table_cell_style)
            
            # Title
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                textColor=colors.HexColor('#003366'),
                spaceAfter=30,
            )
            story.append(Paragraph("NEC Contract Analysis Report", title_style))
            story.append(Spacer(1, 20))
            
            # 1. Executive Summary
            story.append(Paragraph("1. Executive Summary", styles['Heading2']))
            story.append(Spacer(1, 12))
            story.append(Paragraph(
                "This report provides a structured analysis of the NEC Contract Data submitted for review. "
                "Its purpose is to:",
                styles['Normal']
            ))
            story.append(Paragraph("• Summarise extracted contractual information", styles['Normal']))
            story.append(Paragraph("• Identify missing or incomplete content", styles['Normal']))
            story.append(Paragraph("• Highlight contractual risks", styles['Normal']))
            story.append(Paragraph("• Provide a completeness assessment aligned with NEC requirements", styles['Normal']))
            story.append(Spacer(1, 12))
            
            contract_completeness = template_data.get('contract_completeness', {})
            doc_type = contract_completeness.get('document_type', 'unknown')
            is_template = contract_completeness.get('is_template', False)
            filled_pct = contract_completeness.get('filled_percentage', 0.0)
            mandatory_missing = contract_completeness.get('mandatory_missing', 0)
            total_mandatory = contract_completeness.get('total_mandatory', 0)
            
            story.append(Paragraph("<b>Overall Contract Status:</b>", styles['Normal']))
            story.append(Paragraph(f"<b>Document Type:</b> {doc_type}", styles['Normal']))
            story.append(Paragraph(f"<b>Template Detected:</b> {is_template}", styles['Normal']))
            story.append(Paragraph(f"<b>Clause Completeness Score:</b> {filled_pct:.1f}%", styles['Normal']))
            story.append(Paragraph(f"<b>Mandatory Missing Clauses:</b> {mandatory_missing} of {total_mandatory}", styles['Normal']))
            story.append(Spacer(1, 20))
            
            # 2. Contract Overview
            story.append(Paragraph("2. Contract Overview", styles['Heading2']))
            story.append(Spacer(1, 12))
            
            metadata = template_data.get('metadata', {})
            project = template_data.get('project', 'Unknown Project')
            extraction_confidence = metadata.get('extraction_confidence', 0.0)
            if isinstance(extraction_confidence, float):
                extraction_confidence = f"{extraction_confidence * 100:.1f}%" if extraction_confidence <= 1.0 else f"{extraction_confidence:.1f}%"
            
            overview_table_data = [
                [_p('Attribute'), _p('Value')],
                [_p('Project Name'), _p(project, 80)],
                [_p('Total Pages'), _p(str(metadata.get('total_pages', 0)))],
                [_p('File Size'), _p(f"{metadata.get('file_size_bytes', 0)} bytes")],
                [_p('Data Extraction Confidence'), _p(str(extraction_confidence))],
            ]
            # Keep total width <= 6.25" (A4 8.27" - 2" margins) to prevent overflow; repeatRows avoids overlap on split
            overview_table = Table(overview_table_data, colWidths=[2*inch, 4.25*inch], repeatRows=1)
            overview_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#e6eef7')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('TOPPADDING', (0, 0), (-1, -1), 8),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                ('LEFTPADDING', (0, 0), (-1, -1), 6),
                ('RIGHTPADDING', (0, 0), (-1, -1), 6),
                ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                ('GRID', (0, 0), (-1, -1), 1, colors.grey),
                ('FONTSIZE', (0, 1), (-1, -1), 9),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f9f9f9')]),
            ]))
            story.append(overview_table)
            story.append(Spacer(1, 20))
            
            # 3. Key Contract Dates (NEC Clause 3)
            story.append(Paragraph("3. Key Contract Dates (NEC Clause 3)", styles['Heading2']))
            story.append(Spacer(1, 12))
            
            extracted_clauses = template_data.get('extracted_clauses', {})
            
            # Helper function to get clause data
            def get_clause(clause_num: str) -> Dict[str, Any]:
                return extracted_clauses.get(clause_num, {'status': 'missing', 'value': 'Not Provided', 'title': 'N/A'})
            
            # Helper: single Paragraph for "Value: X" (one flowable avoids overlap)
            def _value_para(val):
                if val is None or val == "":
                    val = "Not Provided"
                s = str(val).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                if len(s) > 180:
                    s = s[:177] + "..."
                return Paragraph(s, table_cell_style)
            def _value_line(val):
                s = str(val or "Not Provided").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                if len(s) > 180:
                    s = s[:177] + "..."
                return Paragraph(f"<b>Value:</b> {s}", styles['Normal'])
            
            # 3.1 Starting Date
            clause_31 = get_clause('3.1')
            story.append(Paragraph("<b>3.1 Starting Date</b>", styles['Heading3']))
            story.append(Paragraph(f"<b>Status:</b> {clause_31.get('status', 'missing')}", styles['Normal']))
            story.append(_value_line(clause_31.get('value', 'Not Provided')))
            story.append(Spacer(1, 12))
            
            # 3.2 Possession Date(s)
            clause_32 = get_clause('3.2')
            story.append(Paragraph("<b>3.2 Possession Date(s)</b>", styles['Heading3']))
            story.append(Paragraph(f"<b>Status:</b> {clause_32.get('status', 'missing')}", styles['Normal']))
            story.append(_value_line(clause_32.get('value', 'Not Provided')))
            story.append(Spacer(1, 12))
            
            # 3.3 Completion Date
            clause_33 = get_clause('3.3')
            story.append(Paragraph("<b>3.3 Completion Date</b>", styles['Heading3']))
            story.append(Paragraph(f"<b>Status:</b> {clause_33.get('status', 'missing')}", styles['Normal']))
            story.append(_value_line(clause_33.get('value', 'Not Provided')))
            story.append(Spacer(1, 12))
            
            # Programme Requirements (Clauses 3.5 & 3.6)
            story.append(Paragraph("<b>Programme Requirements (Clauses 3.5 & 3.6)</b>", styles['Heading3']))
            clause_35 = get_clause('3.5')
            clause_36 = get_clause('3.6')
            
            programme_table_data = [
                [_p('Requirement'), _p('Status'), _p('Value')],
                [_p('First Programme Submission'), _p(clause_35.get('status', 'missing')), _p(clause_35.get('value', 'Not Provided'), 120)],
                [_p('Revised Programme Interval'), _p(clause_36.get('status', 'missing')), _p(clause_36.get('value', 'Not Provided'), 120)],
            ]
            
            programme_table = Table(programme_table_data, colWidths=[1.8*inch, 0.9*inch, 2.8*inch], repeatRows=1)
            programme_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#e6eef7')),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('TOPPADDING', (0, 0), (-1, -1), 8),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                ('LEFTPADDING', (0, 0), (-1, -1), 6),
                ('RIGHTPADDING', (0, 0), (-1, -1), 6),
                ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                ('GRID', (0, 0), (-1, -1), 1, colors.grey),
                ('FONTSIZE', (0, 1), (-1, -1), 9),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f9f9f9')]),
            ]))
            story.append(programme_table)
            story.append(Spacer(1, 12))
            
            # Delay Damages (Clause 3.7)
            clause_37 = get_clause('3.7')
            story.append(Paragraph("<b>Delay Damages (Clause 3.7)</b>", styles['Heading3']))
            story.append(Paragraph(f"<b>Status:</b> {clause_37.get('status', 'missing')}", styles['Normal']))
            story.append(_value_line(clause_37.get('value', 'Not Provided')))
            story.append(Spacer(1, 20))
            
            # 4. Testing & Defects (NEC Clause 4)
            story.append(Paragraph("4. Testing & Defects (NEC Clause 4)", styles['Heading2']))
            story.append(Spacer(1, 12))
            
            clause_41 = get_clause('4.1')
            clause_42 = get_clause('4.2')
            clause_43 = get_clause('4.3')
            
            defects_table_data = [
                [_p('Clause'), _p('Description'), _p('Status'), _p('Value')],
                [_p('4.1'), _p('Defects Date'), _p(clause_41.get('status', 'missing')), _p(clause_41.get('value', 'Not Provided'), 80)],
                [_p('4.2'), _p('Defect Correction Period'), _p(clause_42.get('status', 'missing')), _p(clause_42.get('value', 'Not Provided'), 80)],
                [_p('4.3'), _p('Landscaping Maintenance Period'), _p(clause_43.get('status', 'missing')), _p(clause_43.get('value', 'Not Provided'), 80)],
            ]
            
            defects_table = Table(defects_table_data, colWidths=[0.6*inch, 1.4*inch, 0.7*inch, 2.25*inch], repeatRows=1)
            defects_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#e6eef7')),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('TOPPADDING', (0, 0), (-1, -1), 8),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                ('LEFTPADDING', (0, 0), (-1, -1), 6),
                ('RIGHTPADDING', (0, 0), (-1, -1), 6),
                ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                ('GRID', (0, 0), (-1, -1), 1, colors.grey),
                ('FONTSIZE', (0, 1), (-1, -1), 9),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f9f9f9')]),
            ]))
            story.append(defects_table)
            story.append(Spacer(1, 20))
            
            # 5. Payment Terms (NEC Clause 5)
            story.append(Paragraph("5. Payment Terms (NEC Clause 5)", styles['Heading2']))
            story.append(Spacer(1, 12))
            
            clause_52 = get_clause('5.2')
            clause_53 = get_clause('5.3')
            clause_55 = get_clause('5.5')
            clause_56 = get_clause('5.6')
            
            payment_table_data = [
                [_p('Clause'), _p('Term'), _p('Status'), _p('Value')],
                [_p('5.2'), _p('Assessment Interval'), _p(clause_52.get('status', 'missing')), _p(clause_52.get('value', 'Not Provided'), 80)],
                [_p('5.3'), _p('Payment Period'), _p(clause_53.get('status', 'missing')), _p(clause_53.get('value', 'Not Provided'), 80)],
                [_p('5.5'), _p('Retention Percentage'), _p(clause_55.get('status', 'missing')), _p(clause_55.get('value', 'Not Provided'), 80)],
                [_p('5.6'), _p('Bond Amount'), _p(clause_56.get('status', 'missing')), _p(clause_56.get('value', 'Not Provided'), 80)],
            ]
            
            payment_table = Table(payment_table_data, colWidths=[0.6*inch, 1.4*inch, 0.7*inch, 2.25*inch], repeatRows=1)
            payment_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#e6eef7')),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('TOPPADDING', (0, 0), (-1, -1), 8),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                ('LEFTPADDING', (0, 0), (-1, -1), 6),
                ('RIGHTPADDING', (0, 0), (-1, -1), 6),
                ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                ('GRID', (0, 0), (-1, -1), 1, colors.grey),
                ('FONTSIZE', (0, 1), (-1, -1), 9),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f9f9f9')]),
            ]))
            story.append(payment_table)
            story.append(Spacer(1, 20))
            
            # 6. Weather Data (NEC Clause 6)
            story.append(Paragraph("6. Weather Data (NEC Clause 6)", styles['Heading2']))
            story.append(Spacer(1, 12))
            
            clause_61 = get_clause('6.1')
            clause_62 = get_clause('6.2')
            clause_63 = get_clause('6.3')
            
            weather_table_data = [
                [_p('Clause'), _p('Description'), _p('Status'), _p('Value')],
                [_p('6.1'), _p('Weather Recording Location'), _p(clause_61.get('status', 'missing')), _p(clause_61.get('value', 'Not Provided'), 80)],
                [_p('6.2'), _p('Weather Measurement Data'), _p(clause_62.get('status', 'missing')), _p(clause_62.get('value', 'Not Provided'), 80)],
                [_p('6.3'), _p('Historical Weather Records Source'), _p(clause_63.get('status', 'missing')), _p(clause_63.get('value', 'Not Provided'), 80)],
            ]
            
            weather_table = Table(weather_table_data, colWidths=[0.6*inch, 1.4*inch, 0.7*inch, 2.25*inch], repeatRows=1)
            weather_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#e6eef7')),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('TOPPADDING', (0, 0), (-1, -1), 8),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                ('LEFTPADDING', (0, 0), (-1, -1), 6),
                ('RIGHTPADDING', (0, 0), (-1, -1), 6),
                ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                ('GRID', (0, 0), (-1, -1), 1, colors.grey),
                ('FONTSIZE', (0, 1), (-1, -1), 9),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f9f9f9')]),
            ]))
            story.append(weather_table)
            story.append(Spacer(1, 20))
            
            # 7. Clause Completeness Assessment
            story.append(Paragraph("7. Clause Completeness Assessment", styles['Heading2']))
            story.append(Spacer(1, 12))
            
            story.append(Paragraph("<b>7.1 Summary</b>", styles['Heading3']))
            mandatory_filled = contract_completeness.get('mandatory_filled', 0)
            mandatory_blank = contract_completeness.get('mandatory_blank', 0)
            
            story.append(Paragraph(f"• <b>Total Mandatory Clauses:</b> {total_mandatory}", styles['Normal']))
            story.append(Paragraph(f"• <b>Filled:</b> {mandatory_filled}", styles['Normal']))
            story.append(Paragraph(f"• <b>Blank:</b> {mandatory_blank}", styles['Normal']))
            story.append(Paragraph(f"• <b>Missing:</b> {mandatory_missing}", styles['Normal']))
            story.append(Spacer(1, 12))
            
            story.append(Paragraph("<b>7.2 Completeness Table</b>", styles['Heading3']))
            
            # Build completeness table (Paragraph cells so long titles wrap)
            completeness_table_data = [[_p('Clause'), _p('Title'), _p('Status')]]
            for clause_num in sorted(extracted_clauses.keys()):
                clause = extracted_clauses[clause_num]
                completeness_table_data.append([
                    _p(clause_num),
                    _p(clause.get('title', 'N/A'), 150),
                    _p(clause.get('status', 'unknown')),
                ])
            
            if len(completeness_table_data) > 1:
                completeness_table = Table(completeness_table_data, colWidths=[0.75*inch, 2.75*inch, 1.75*inch], repeatRows=1)
                completeness_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#e6eef7')),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 10),
                    ('TOPPADDING', (0, 0), (-1, -1), 8),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                    ('LEFTPADDING', (0, 0), (-1, -1), 6),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 6),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                    ('GRID', (0, 0), (-1, -1), 1, colors.grey),
                    ('FONTSIZE', (0, 1), (-1, -1), 9),
                    ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f9f9f9')]),
                ]))
                story.append(completeness_table)
            story.append(Spacer(1, 20))
            
            # 8. Risk Analysis
            story.append(Paragraph("8. Risk Analysis", styles['Heading2']))
            story.append(Spacer(1, 12))
            
            story.append(Paragraph("<b>🔴 Critical Risks</b>", styles['Heading3']))
            story.append(Paragraph("• Missing key dates prevents programme development", styles['Normal']))
            story.append(Paragraph("• Missing payment terms creates dispute risk", styles['Normal']))
            story.append(Paragraph("• Missing weather station details affects compensation event compliance", styles['Normal']))
            story.append(Paragraph("• Missing defects data weakens handover obligations", styles['Normal']))
            story.append(Spacer(1, 12))
            
            story.append(Paragraph("<b>🟠 Moderate Risks</b>", styles['Heading3']))
            story.append(Paragraph("• Programme submission and revision rules missing", styles['Normal']))
            story.append(Paragraph("• Defect correction period unclear", styles['Normal']))
            story.append(Spacer(1, 12))
            
            story.append(Paragraph("<b>🟡 Observations</b>", styles['Heading3']))
            if is_template:
                story.append(Paragraph("• Contract appears to be a <b>template</b>, not a completed contract", styles['Normal']))
            story.append(Paragraph("• Commercial and scheduling data are absent", styles['Normal']))
            story.append(Spacer(1, 20))
            
            # 9. Recommendations
            story.append(Paragraph("9. Recommendations", styles['Heading2']))
            story.append(Spacer(1, 12))
            recommendations = [
                "Populate all mandatory programme clauses (3.1–3.7)",
                "Confirm payment cycle and assessment rules",
                "Provide weather station details before execution",
                "Complete Defects and Maintenance periods",
                "Ensure key dates are provided to support programme creation",
            ]
            for rec in recommendations:
                story.append(Paragraph(f"• {rec}", styles['Normal']))
            story.append(Spacer(1, 20))
            
            # 10. Appendices
            story.append(Paragraph("10. Appendices", styles['Heading2']))
            story.append(Spacer(1, 12))
            story.append(Paragraph("• <b>Appendix A:</b> Raw JSON Extract", styles['Normal']))
            story.append(Paragraph("• <b>Appendix B:</b> Clause-by-Clause Log", styles['Normal']))
            story.append(Paragraph("• <b>Appendix C:</b> Programme Requirements Summary", styles['Normal']))
            
            # Build PDF
            doc.build(story)
            pdf_bytes = buffer.getvalue()
            buffer.close()
            
            # Save to file if output_path provided
            if output_path:
                output_path = str(Path(output_path).resolve())
                Path(output_path).parent.mkdir(parents=True, exist_ok=True)
                with open(output_path, 'wb') as f:
                    f.write(pdf_bytes)
                print(f"[ReportGenerator] PDF generated: {output_path}")
                return output_path
            
            print(f"[ReportGenerator] PDF generated ({len(pdf_bytes)} bytes)")
            return pdf_bytes
            
        except Exception as pdf_error:
            print(f"[ReportGenerator] PDF generation failed: {pdf_error}, attempting DOCX fallback...")
            if DOCX_AVAILABLE:
                return self.generate_docx(json_data, output_path)
            else:
                raise Exception(f"PDF generation failed and DOCX fallback unavailable: {pdf_error}")
    
    def generate_docx(
        self,
        json_data: Dict[str, Any],
        output_path: Optional[str] = None,
        report_options: Optional[Dict[str, Any]] = None,
    ) -> Union[bytes, str]:
        """
        Generates a professional DOCX report from JSON data.
        If input is validation JSON (alignment + contract_summary), produces Programme Validation Report (Sections A–G).
        """
        if "alignment" in json_data and "contract_summary" in json_data:
            return self._generate_validation_docx(json_data, output_path, report_options=report_options)
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx is required for DOCX generation. "
                "Install it with: pip install python-docx"
            )
        
        # Prepare data for template rendering (contract analysis)
        template_data = self._prepare_template_data(json_data)
        
        # Create document
        doc = Document()
        
        # Title
        title = doc.add_heading('NEC Contract Analysis Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 1. Executive Summary
        doc.add_heading('1. Executive Summary', 1)
        doc.add_paragraph(
            'This report provides a structured analysis of the NEC Contract Data submitted for review. '
            'Its purpose is to:'
        )
        doc.add_paragraph('• Summarise extracted contractual information', style='List Bullet')
        doc.add_paragraph('• Identify missing or incomplete content', style='List Bullet')
        doc.add_paragraph('• Highlight contractual risks', style='List Bullet')
        doc.add_paragraph('• Provide a completeness assessment aligned with NEC requirements', style='List Bullet')
        
        contract_completeness = template_data.get('contract_completeness', {})
        doc_type = contract_completeness.get('document_type', 'unknown')
        is_template = contract_completeness.get('is_template', False)
        filled_pct = contract_completeness.get('filled_percentage', 0.0)
        mandatory_missing = contract_completeness.get('mandatory_missing', 0)
        total_mandatory = contract_completeness.get('total_mandatory', 0)
        
        p = doc.add_paragraph()
        p.add_run('Overall Contract Status:').bold = True
        doc.add_paragraph(f'Document Type: {doc_type}')
        doc.add_paragraph(f'Template Detected: {is_template}')
        doc.add_paragraph(f'Clause Completeness Score: {filled_pct:.1f}%')
        doc.add_paragraph(f'Mandatory Missing Clauses: {mandatory_missing} of {total_mandatory}')
        
        # 2. Contract Overview
        doc.add_heading('2. Contract Overview', 1)
        metadata = template_data.get('metadata', {})
        project = template_data.get('project', 'Unknown Project')
        extraction_confidence = metadata.get('extraction_confidence', 0.0)
        if isinstance(extraction_confidence, float):
            extraction_confidence = f"{extraction_confidence * 100:.1f}%" if extraction_confidence <= 1.0 else f"{extraction_confidence:.1f}%"
        
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Light Grid Accent 1'
        table.rows[0].cells[0].text = 'Attribute'
        table.rows[0].cells[1].text = 'Value'
        table.rows[1].cells[0].text = 'Project Name'
        table.rows[1].cells[1].text = project
        table.rows[2].cells[0].text = 'Total Pages'
        table.rows[2].cells[1].text = str(metadata.get('total_pages', 0))
        table.rows[3].cells[0].text = 'File Size'
        table.rows[3].cells[1].text = f"{metadata.get('file_size_bytes', 0)} bytes"
        table.rows[4].cells[0].text = 'Data Extraction Confidence'
        table.rows[4].cells[1].text = str(extraction_confidence)
        
        # 3. Key Contract Dates
        doc.add_heading('3. Key Contract Dates (NEC Clause 3)', 1)
        extracted_clauses = template_data.get('extracted_clauses', {})
        
        def get_clause(clause_num: str) -> Dict[str, Any]:
            return extracted_clauses.get(clause_num, {'status': 'missing', 'value': 'Not Provided', 'title': 'N/A'})
        
        clause_31 = get_clause('3.1')
        doc.add_heading('3.1 Starting Date', 2)
        doc.add_paragraph(f'Status: {clause_31.get("status", "missing")}')
        doc.add_paragraph(f'Value: {clause_31.get("value", "Not Provided")}')
        
        clause_32 = get_clause('3.2')
        doc.add_heading('3.2 Possession Date(s)', 2)
        doc.add_paragraph(f'Status: {clause_32.get("status", "missing")}')
        doc.add_paragraph(f'Value: {clause_32.get("value", "Not Provided")}')
        
        clause_33 = get_clause('3.3')
        doc.add_heading('3.3 Completion Date', 2)
        doc.add_paragraph(f'Status: {clause_33.get("status", "missing")}')
        doc.add_paragraph(f'Value: {clause_33.get("value", "Not Provided")}')
        
        doc.add_heading('Programme Requirements (Clauses 3.5 & 3.6)', 2)
        clause_35 = get_clause('3.5')
        clause_36 = get_clause('3.6')
        prog_table = doc.add_table(rows=3, cols=3)
        prog_table.style = 'Light Grid Accent 1'
        prog_table.rows[0].cells[0].text = 'Requirement'
        prog_table.rows[0].cells[1].text = 'Status'
        prog_table.rows[0].cells[2].text = 'Value'
        prog_table.rows[1].cells[0].text = 'First Programme Submission'
        prog_table.rows[1].cells[1].text = clause_35.get('status', 'missing')
        prog_table.rows[1].cells[2].text = clause_35.get('value', 'Not Provided')
        prog_table.rows[2].cells[0].text = 'Revised Programme Interval'
        prog_table.rows[2].cells[1].text = clause_36.get('status', 'missing')
        prog_table.rows[2].cells[2].text = clause_36.get('value', 'Not Provided')
        
        clause_37 = get_clause('3.7')
        doc.add_heading('Delay Damages (Clause 3.7)', 2)
        doc.add_paragraph(f'Status: {clause_37.get("status", "missing")}')
        doc.add_paragraph(f'Value: {clause_37.get("value", "Not Provided")}')
        
        # 4. Testing & Defects
        doc.add_heading('4. Testing & Defects (NEC Clause 4)', 1)
        clause_41 = get_clause('4.1')
        clause_42 = get_clause('4.2')
        clause_43 = get_clause('4.3')
        
        defects_table = doc.add_table(rows=4, cols=4)
        defects_table.style = 'Light Grid Accent 1'
        defects_table.rows[0].cells[0].text = 'Clause'
        defects_table.rows[0].cells[1].text = 'Description'
        defects_table.rows[0].cells[2].text = 'Status'
        defects_table.rows[0].cells[3].text = 'Value'
        defects_table.rows[1].cells[0].text = '4.1'
        defects_table.rows[1].cells[1].text = 'Defects Date'
        defects_table.rows[1].cells[2].text = clause_41.get('status', 'missing')
        defects_table.rows[1].cells[3].text = clause_41.get('value', 'Not Provided')
        defects_table.rows[2].cells[0].text = '4.2'
        defects_table.rows[2].cells[1].text = 'Defect Correction Period'
        defects_table.rows[2].cells[2].text = clause_42.get('status', 'missing')
        defects_table.rows[2].cells[3].text = clause_42.get('value', 'Not Provided')
        defects_table.rows[3].cells[0].text = '4.3'
        defects_table.rows[3].cells[1].text = 'Landscaping Maintenance Period'
        defects_table.rows[3].cells[2].text = clause_43.get('status', 'missing')
        defects_table.rows[3].cells[3].text = clause_43.get('value', 'Not Provided')
        
        # 5. Payment Terms
        doc.add_heading('5. Payment Terms (NEC Clause 5)', 1)
        clause_52 = get_clause('5.2')
        clause_53 = get_clause('5.3')
        clause_55 = get_clause('5.5')
        clause_56 = get_clause('5.6')
        
        payment_table = doc.add_table(rows=5, cols=4)
        payment_table.style = 'Light Grid Accent 1'
        payment_table.rows[0].cells[0].text = 'Clause'
        payment_table.rows[0].cells[1].text = 'Term'
        payment_table.rows[0].cells[2].text = 'Status'
        payment_table.rows[0].cells[3].text = 'Value'
        payment_table.rows[1].cells[0].text = '5.2'
        payment_table.rows[1].cells[1].text = 'Assessment Interval'
        payment_table.rows[1].cells[2].text = clause_52.get('status', 'missing')
        payment_table.rows[1].cells[3].text = clause_52.get('value', 'Not Provided')
        payment_table.rows[2].cells[0].text = '5.3'
        payment_table.rows[2].cells[1].text = 'Payment Period'
        payment_table.rows[2].cells[2].text = clause_53.get('status', 'missing')
        payment_table.rows[2].cells[3].text = clause_53.get('value', 'Not Provided')
        payment_table.rows[3].cells[0].text = '5.5'
        payment_table.rows[3].cells[1].text = 'Retention Percentage'
        payment_table.rows[3].cells[2].text = clause_55.get('status', 'missing')
        payment_table.rows[3].cells[3].text = clause_55.get('value', 'Not Provided')
        payment_table.rows[4].cells[0].text = '5.6'
        payment_table.rows[4].cells[1].text = 'Bond Amount'
        payment_table.rows[4].cells[2].text = clause_56.get('status', 'missing')
        payment_table.rows[4].cells[3].text = clause_56.get('value', 'Not Provided')
        
        # 6. Weather Data
        doc.add_heading('6. Weather Data (NEC Clause 6)', 1)
        clause_61 = get_clause('6.1')
        clause_62 = get_clause('6.2')
        clause_63 = get_clause('6.3')
        
        weather_table = doc.add_table(rows=4, cols=4)
        weather_table.style = 'Light Grid Accent 1'
        weather_table.rows[0].cells[0].text = 'Clause'
        weather_table.rows[0].cells[1].text = 'Description'
        weather_table.rows[0].cells[2].text = 'Status'
        weather_table.rows[0].cells[3].text = 'Value'
        weather_table.rows[1].cells[0].text = '6.1'
        weather_table.rows[1].cells[1].text = 'Weather Recording Location'
        weather_table.rows[1].cells[2].text = clause_61.get('status', 'missing')
        weather_table.rows[1].cells[3].text = clause_61.get('value', 'Not Provided')
        weather_table.rows[2].cells[0].text = '6.2'
        weather_table.rows[2].cells[1].text = 'Weather Measurement Data'
        weather_table.rows[2].cells[2].text = clause_62.get('status', 'missing')
        weather_table.rows[2].cells[3].text = clause_62.get('value', 'Not Provided')
        weather_table.rows[3].cells[0].text = '6.3'
        weather_table.rows[3].cells[1].text = 'Historical Weather Records Source'
        weather_table.rows[3].cells[2].text = clause_63.get('status', 'missing')
        weather_table.rows[3].cells[3].text = clause_63.get('value', 'Not Provided')
        
        # 7. Clause Completeness Assessment
        doc.add_heading('7. Clause Completeness Assessment', 1)
        doc.add_heading('7.1 Summary', 2)
        mandatory_filled = contract_completeness.get('mandatory_filled', 0)
        mandatory_blank = contract_completeness.get('mandatory_blank', 0)
        
        doc.add_paragraph(f'• Total Mandatory Clauses: {total_mandatory}', style='List Bullet')
        doc.add_paragraph(f'• Filled: {mandatory_filled}', style='List Bullet')
        doc.add_paragraph(f'• Blank: {mandatory_blank}', style='List Bullet')
        doc.add_paragraph(f'• Missing: {mandatory_missing}', style='List Bullet')
        
        doc.add_heading('7.2 Completeness Table', 2)
        completeness_table = doc.add_table(rows=1, cols=3)
        completeness_table.style = 'Light Grid Accent 1'
        completeness_table.rows[0].cells[0].text = 'Clause'
        completeness_table.rows[0].cells[1].text = 'Title'
        completeness_table.rows[0].cells[2].text = 'Status'
        
        for clause_num in sorted(extracted_clauses.keys()):
            clause = extracted_clauses[clause_num]
            row = completeness_table.add_row()
            row.cells[0].text = clause_num
            row.cells[1].text = clause.get('title', 'N/A')
            row.cells[2].text = clause.get('status', 'unknown')
        
        # 8. Risk Analysis
        doc.add_heading('8. Risk Analysis', 1)
        doc.add_heading('🔴 Critical Risks', 2)
        doc.add_paragraph('• Missing key dates prevents programme development', style='List Bullet')
        doc.add_paragraph('• Missing payment terms creates dispute risk', style='List Bullet')
        doc.add_paragraph('• Missing weather station details affects compensation event compliance', style='List Bullet')
        doc.add_paragraph('• Missing defects data weakens handover obligations', style='List Bullet')
        
        doc.add_heading('🟠 Moderate Risks', 2)
        doc.add_paragraph('• Programme submission and revision rules missing', style='List Bullet')
        doc.add_paragraph('• Defect correction period unclear', style='List Bullet')
        
        doc.add_heading('🟡 Observations', 2)
        if is_template:
            p = doc.add_paragraph('• Contract appears to be a ', style='List Bullet')
            p.add_run('template').bold = True
            p.add_run(', not a completed contract')
        doc.add_paragraph('• Commercial and scheduling data are absent', style='List Bullet')
        
        # 9. Recommendations
        doc.add_heading('9. Recommendations', 1)
        recommendations = [
            'Populate all mandatory programme clauses (3.1–3.7)',
            'Confirm payment cycle and assessment rules',
            'Provide weather station details before execution',
            'Complete Defects and Maintenance periods',
            'Ensure key dates are provided to support programme creation',
        ]
        for rec in recommendations:
            doc.add_paragraph(f'• {rec}', style='List Bullet')
        
        # 10. Appendices
        doc.add_heading('10. Appendices', 1)
        doc.add_paragraph('• Appendix A: Raw JSON Extract', style='List Bullet')
        doc.add_paragraph('• Appendix B: Clause-by-Clause Log', style='List Bullet')
        doc.add_paragraph('• Appendix C: Programme Requirements Summary', style='List Bullet')
        
        # Save to file or return bytes
        if output_path:
            output_path = str(Path(output_path).resolve())
            Path(output_path).parent.mkdir(parents=True, exist_ok=True)
            doc.save(output_path)
            print(f"[ReportGenerator] DOCX generated: {output_path}")
            return output_path
        else:
            buffer = BytesIO()
            doc.save(buffer)
            docx_bytes = buffer.getvalue()
            buffer.close()
            print(f"[ReportGenerator] DOCX generated ({len(docx_bytes)} bytes)")
            return docx_bytes
    
    def _generate_validation_pdf(
        self,
        json_data: Dict[str, Any],
        output_path: Optional[str] = None,
        report_options: Optional[Dict[str, Any]] = None,
    ) -> Union[bytes, str]:
        """Generate Programme Validation Report (Sections A–G) as PDF. Flow-based layout only (Platypus)."""
        from app.reporting.validation_report_builder import build_validation_report
        opts = report_options or {}
        redact_activities = opts.get("confidentiality_mode") is True
        user_name = (opts.get("user_name") or "").strip()
        logo_url = (opts.get("organisation_logo_url") or "").strip()

        def _redact(s: Any) -> str:
            if not redact_activities:
                return str(s).strip() if s is not None and str(s).strip() else "—"
            return "—"

        report = build_validation_report(json_data)
        report = report if isinstance(report, dict) else {}
        if not REPORTLAB_AVAILABLE:
            raise ImportError("ReportLab not available")
        histogram_temp_path = None  # for cleanup after build (Schedule Float image)
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=72)
        story = []
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle('ValidationTitle', parent=styles['Heading1'], fontSize=22, textColor=colors.HexColor('#003366'), spaceAfter=24)
        section_heading_style = ParagraphStyle('ValidationSectionHeading', parent=styles['Heading2'], fontSize=13, fontName='Helvetica-Bold', spaceAfter=6, spaceBefore=4)
        # Table cell style: wrap text in cells to avoid overlap; no absolute positioning
        table_cell_style = ParagraphStyle(
            "ValidationTableCell",
            parent=styles["Normal"],
            fontSize=7,
            leading=11,
            spaceBefore=3,
            spaceAfter=3,
            wordWrap="CJK",
        )

        def _p(s: str, max_len: Optional[int] = None) -> "Paragraph":
            if s is None:
                s = "—"
            s = str(s).strip()
            if max_len and len(s) > max_len:
                s = s[: max_len - 3] + "..."
            s = s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;") or "—"
            return Paragraph(s, table_cell_style)

        def _validation_table_style():
            return TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#e6eef7')),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 9),
                ('TOPPADDING', (0, 0), (-1, -1), 8),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                ('LEFTPADDING', (0, 0), (-1, -1), 6),
                ('RIGHTPADDING', (0, 0), (-1, -1), 6),
                ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('FONTSIZE', (0, 1), (-1, -1), 8),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f9f9f9')]),
            ])

        def _sec(k: str) -> dict:
            """Return section as dict; if value is string, wrap as {_narrative: s} for direct render."""
            val = report.get(k)
            if isinstance(val, str):
                return {"_narrative": val}
            return _safe_dict(val)

        def _render_section(val: Any, styles: Any) -> None:
            """Append content: string→paragraph, dict with _narrative→paragraph, else no-op."""
            if isinstance(val, str) and val:
                story.append(Paragraph(val.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"), styles["Normal"]))
            elif isinstance(val, dict) and val.get("_narrative"):
                story.append(Paragraph(str(val["_narrative"]).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"), styles["Normal"]))

        meta = _sec("metadata")
        if not isinstance(meta, dict):
            meta = {}

        float_stats = _extract_float_stats_from_validation(json_data)

        # Plan Analytics header and metadata (branding; no change to acceptability logic)
        header_style = ParagraphStyle('PlanAnalyticsHeader', parent=styles['Normal'], fontSize=10, textColor=colors.HexColor('#64748b'), spaceAfter=4)
        story.append(Paragraph("Plan Analytics", header_style))
        run_id = json_data.get("run_id")
        submission_id = str(run_id) if run_id is not None else "—"
        gen_ts = meta.get("validation_timestamp") or datetime.now().isoformat()
        story.append(Paragraph(f"<b>Submission ID:</b> {submission_id}  |  <b>Generated:</b> {gen_ts[:19] if len(gen_ts) > 19 else gen_ts}", styles['Normal']))
        story.append(Spacer(1, 12))

        # Section A — Executive Summary
        story.append(Paragraph("Programme Validation Report", title_style))
        story.append(Paragraph(f"<i>Contract: {meta.get('contract_file', '—')}  |  Programme: {meta.get('xer_file', '—')}</i>", styles['Normal']))
        story.append(Spacer(1, 14))
        story.append(Paragraph("Section A — Executive Summary", section_heading_style))
        story.append(Spacer(1, 8))
        a = _sec("section_a_executive_summary")
        if a.get("_narrative"):
            story.append(Paragraph(str(a["_narrative"]).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;") or "—", styles["Normal"]))
        elif a.get("decision_heading"):
            story.append(Paragraph(f"<b>{a.get('decision_heading')}</b>", styles['Normal']))
            story.append(Spacer(1, 6))
        if a.get("decision_detail"):
            story.append(Paragraph(a.get("decision_detail"), styles['Normal']))
        if (a.get("decision_heading") or a.get("decision_detail")):
            if float_stats:
                risk_level = _delivery_risk_profile(float_stats["critical_pct"], float_stats["under_40_pct"])
                story.append(Paragraph(f"Delivery risk profile: {risk_level}", styles['Normal']))
                story.append(Paragraph(_delivery_risk_explanation(risk_level), styles['Normal']))
            story.append(Paragraph(
                "This means the programme meets NEC Clause 31 submission requirements at this stage based on demonstrated alignment of mandatory obligations.",
                styles['Normal']
            ))
        if a.get("programme_stage_context"):
            story.append(Spacer(1, 6))
            story.append(Paragraph(f"<i>{a.get('programme_stage_context')}</i>", styles['Normal']))
        if a.get("headline_reason"):
            story.append(Spacer(1, 6))
            story.append(Paragraph(f"<i>Main reason:</i> {a.get('headline_reason')}", styles['Normal']))
        if a.get("quality_summary"):
            story.append(Spacer(1, 6))
            story.append(Paragraph(a.get("quality_summary"), styles['Normal']))
        if a.get("quality_detail"):
            story.append(Paragraph(a.get("quality_detail"), styles['Normal']))
        if a.get("reassurance"):
            story.append(Spacer(1, 6))
            story.append(Paragraph(a.get("reassurance"), styles['Normal']))
        if a.get("executive_summary_text"):
            story.append(Spacer(1, 10))
            story.append(Paragraph(a.get("executive_summary_text"), styles['Normal']))
        story.append(Spacer(1, 14))

        # Section B — What determined the decision
        story.append(Paragraph("Section B — What determined the decision", section_heading_style))
        story.append(Spacer(1, 8))
        b = _sec("section_b_what_determined_outcome")
        if b.get("section_intro"):
            story.append(Paragraph(b.get("section_intro"), styles['Normal']))
            story.append(Spacer(1, 8))
        attention_items = b.get("items_requiring_attention", []) or []
        reassurance_items = b.get("items_in_good_order", []) or []
        if attention_items:
            story.append(Paragraph("<b>Items needing action</b>", styles['Heading3']))
            table_data = [[_p("Contract check"), _p("What we saw"), _p("What needs to happen")]]
            for row in attention_items[:20]:
                r = _safe_row(row)
                table_data.append([
                    _p(r.get("requirement", ""), 80),
                    _p(_presentation_text(r.get("observation", "")), 120),
                    _p(r.get("action", ""), 120),
                ])
            t = Table(table_data, colWidths=[1.5*inch, 2.3*inch, 2.3*inch], repeatRows=1)
            t.setStyle(_validation_table_style())
            story.append(t)
            story.append(Spacer(1, 12))
        if reassurance_items:
            story.append(Paragraph("<b>Items already in good order</b>", styles['Heading3']))
            for row in reassurance_items[:15]:
                r = _safe_row(row)
                story.append(Paragraph(f"• {r.get('message', '')}", styles['Normal']))
            story.append(Spacer(1, 12))
        if b.get("alternative_interpretation"):
            story.append(Paragraph(f"<i>{b.get('alternative_interpretation')}</i>", styles['Normal']))
            story.append(Spacer(1, 8))

        # Section — Scope and constraints coverage
        scope_sec = _sec("section_scope_contract_alignment")
        story.append(Paragraph("How scope and constraints are assessed", section_heading_style))
        story.append(Spacer(1, 8))
        story.append(Paragraph(
            "This section evaluates whether contractual scope and constraints are clearly represented in the programme through activities, sequencing, and access logic.",
            styles['Normal']
        ))
        story.append(Spacer(1, 10))
        overall_conf = scope_sec.get("overall_confidence")
        if overall_conf is None or (isinstance(overall_conf, str) and not overall_conf.strip()):
            overall_conf = _derive_overall_confidence(json_data)
        else:
            overall_conf = str(overall_conf).strip()
        if overall_conf not in ("High", "Medium", "Low"):
            overall_conf = _derive_overall_confidence(json_data)
        story.append(Paragraph(f"<b>Overall confidence: {overall_conf}</b>", styles['Normal']))
        story.append(Paragraph(
            "Confidence levels are assessed as follows:",
            styles['Normal']
        ))
        story.append(Paragraph("• High — Direct evidence of obligation alignment and sequencing", styles['Normal']))
        story.append(Paragraph("• Medium — Alignment demonstrated with limited assumptions", styles['Normal']))
        story.append(Paragraph("• Low — Material reliance on interpretation or missing sequencing evidence", styles['Normal']))
        story.append(Spacer(1, 6))
        scope_rows_raw = scope_sec.get("scope_rows", []) or []
        scope_rows = []
        for row in scope_rows_raw:
            if isinstance(row, dict):
                scope_rows.append(row)
            else:
                scope_rows.append({
                    "contract_scope": str(row),
                    "programme_activities": [],
                    "representation_status": "Not specified",
                    "notes": "",
                })
        if scope_rows:
            story.append(Paragraph("<b>Contract scope evidence</b>", styles['Heading3']))
            story.append(Spacer(1, 6))
            table_data = [
                [_p("Contract scope item"), _p("Programme activities"), _p("Coverage strength"), _p("Notes"), _p("Confidence")]
            ]
            for row in scope_rows[:25]:
                r = _safe_row(row)
                pa = r.get("programme_activities", [])
                pa = pa if isinstance(pa, list) else ([pa] if pa is not None else [])
                programme_evidence = "\n".join(f"• {x}" for x in pa if x) or "—"
                table_data.append([
                    _p(r.get("contract_scope", ""), 90),
                    _p(programme_evidence, 70),
                    _p(r.get("representation_status", ""), 60),
                    _p(_presentation_text(r.get("notes", "")), 100),
                    _p(_confidence_display(r.get("confidence_band", "")), 38),
                ])
            t = Table(table_data, colWidths=[2.1*inch, 1.6*inch, 1.4*inch, 2.1*inch, 1.0*inch], repeatRows=1)
            t.setStyle(_validation_table_style())
            story.append(t)
            activity_load_notes = scope_sec.get("activity_load_notes") or []
            if activity_load_notes:
                story.append(Spacer(1, 8))
                story.append(Paragraph("<b>Activity load transparency</b>", styles['Heading3']))
                for note in activity_load_notes[:10]:
                    story.append(Paragraph(f"• {note}", styles['Normal']))
            story.append(Spacer(1, 10))
        # Evidence-first: list evidenced, explicit assumption, not-represented-but-mandatory (governance section omitted)
        obligations_evidenced_list = scope_sec.get("obligations_evidenced_list") or []
        obligations_explicit_assumption_list = scope_sec.get("obligations_explicit_assumption_list") or []
        obligations_not_represented_but_mandatory_list = scope_sec.get("obligations_not_represented_but_mandatory_list") or []
        if obligations_evidenced_list:
            story.append(Paragraph("<b>Obligations evidenced by programme</b>", styles['Heading3']))
            story.append(Paragraph("The following obligations are evidenced by programme activities or constraints.", styles['Normal']))
            for ob in obligations_evidenced_list[:30]:
                o = ob if isinstance(ob, dict) else {}
                story.append(Paragraph(f"• [{o.get('id', '')}] {o.get('text', '')}", styles['Normal']))
            story.append(Spacer(1, 6))
        if obligations_explicit_assumption_list:
            story.append(Paragraph("<b>Obligations with explicit assumption</b>", styles['Heading3']))
            story.append(Paragraph("The following are aligned via explicit assumption (client responsibility or out of scope at this stage).", styles['Normal']))
            for ob in obligations_explicit_assumption_list[:30]:
                o = ob if isinstance(ob, dict) else {}
                story.append(Paragraph(f"• [{o.get('id', '')}] {o.get('text', '')}", styles['Normal']))
                if o.get("exemption_reason"):
                    story.append(Paragraph(f"  <i>{o.get('exemption_reason')}</i>", styles['Normal']))
            story.append(Spacer(1, 6))
        obligations_covered_by_later_submission_list = scope_sec.get("obligations_covered_by_later_submission_list") or []
        if obligations_covered_by_later_submission_list:
            story.append(Paragraph("<b>Obligations assumed to be covered by later submission (non-blocking advisory only)</b>", styles['Heading3']))
            story.append(Paragraph("The following have only a 'covered by later submission' assumption. They do not set aligned and must never justify acceptance.", styles['Normal']))
            for ob in obligations_covered_by_later_submission_list[:30]:
                o = ob if isinstance(ob, dict) else {}
                story.append(Paragraph(f"• [{o.get('id', '')}] {o.get('text', '')}", styles['Normal']))
            story.append(Spacer(1, 6))
        if obligations_not_represented_but_mandatory_list:
            story.append(Paragraph("<b>Not represented but mandatory</b>", styles['Heading3']))
            story.append(Paragraph("The following mandatory obligations are not represented in the programme and block acceptability. They require either explicit programme activities or an explicit assumption.", styles['Normal']))
            for ob in obligations_not_represented_but_mandatory_list[:30]:
                o = ob if isinstance(ob, dict) else {}
                story.append(Paragraph(f"• [{o.get('id', '')}] {o.get('text', '')}", styles['Normal']))
            story.append(Spacer(1, 6))
        # Requires governance / future submission section removed from report output.
        if scope_sec.get("constraint_summary"):
            story.append(Paragraph(scope_sec.get("constraint_summary"), styles['Normal']))
            story.append(Spacer(1, 6))
        constraint_rows_raw = scope_sec.get("constraint_rows", []) or []
        constraint_rows = []
        for row in constraint_rows_raw:
            if isinstance(row, dict):
                constraint_rows.append(row)
            else:
                constraint_rows.append({
                    "constraint": str(row),
                    "programme_evidence": "",
                    "handling": "Implicitly managed",
                })
        if constraint_rows:
            story.append(Paragraph("<b>Constraint coverage</b>", styles['Heading3']))
            story.append(Spacer(1, 6))
            table_data = [
                [_p("Constraint"), _p("Programme evidence"), _p("Handling"), _p("Confidence")]
            ]
            for row in constraint_rows[:25]:
                r = _safe_row(row)
                table_data.append([
                    _p(r.get("constraint", ""), 90),
                    _p(r.get("programme_evidence", ""), 80),
                    _p(r.get("handling", ""), 50),
                    _p(_confidence_display(r.get("confidence_band", "")), 38),
                ])
            t = Table(table_data, colWidths=[2.5*inch, 2.4*inch, 1.0*inch, 1.0*inch], repeatRows=1)
            t.setStyle(_validation_table_style())
            story.append(t)
            story.append(Spacer(1, 10))
        if scope_sec.get("acceptability_clarification"):
            story.append(Paragraph(scope_sec.get("acceptability_clarification"), styles['Normal']))
            story.append(Spacer(1, 6))
        if scope_sec.get("reassurance"):
            story.append(Paragraph(scope_sec.get("reassurance"), styles['Normal']))
            story.append(Spacer(1, 14))

        # Section C — Programme vs Contract Dates
        story.append(Paragraph("Section C — Programme vs Contract Dates", section_heading_style))
        story.append(Spacer(1, 8))
        c = _sec("section_c_dates")
        if c.get("section_intro"):
            story.append(Paragraph(c.get("section_intro"), styles['Normal']))
            story.append(Spacer(1, 8))
        rows = c.get("rows", []) or []
        table_data = [[_p("Item"), _p("Contract"), _p("Programme"), _p("Notes")]]
        for r in rows:
            r = _safe_row(r)
            table_data.append([
                _p(r.get("item", ""), 50),
                _p(r.get("contract", ""), 36),
                _p(r.get("programme", ""), 36),
                _p(_presentation_text(r.get("notes", "")), 70),
            ])
        t = Table(table_data, colWidths=[1.4*inch, 1.3*inch, 1.3*inch, 2*inch], repeatRows=1)
        t.setStyle(_validation_table_style())
        story.append(t)
        story.append(Spacer(1, 8))
        variance = _presentation_text(c.get('variance_explanation') or '')[:500]
        story.append(Paragraph(f"<i>Variance explanation:</i> {variance}", styles['Normal']))
        story.append(Spacer(1, 14))

        # Section D (Required activities) removed for clarity — not included in PDF output.

        # Section E — Programme Quality & Realism (Advisory)
        story.append(Paragraph("Section E — Programme confidence and observations", section_heading_style))
        story.append(Spacer(1, 8))
        e = _sec("section_e_quality_realism_advisory")
        if e.get("section_intro"):
            story.append(Paragraph(e.get("section_intro"), styles['Normal']))
            story.append(Spacer(1, 8))
        for obs in e.get("observations", [])[:10]:
            story.append(Paragraph(f"• {obs}", styles['Normal']))
        if e.get("supportive_notes"):
            story.append(Spacer(1, 6))
            story.append(Paragraph("<i>Additional notes:</i>", styles['Normal']))
            for note in e.get("supportive_notes", [])[:5]:
                story.append(Paragraph(f"• {note}", styles['Normal']))
        if e.get("quality_commentary"):
            story.append(Spacer(1, 6))
            story.append(Paragraph(e.get("quality_commentary"), styles['Normal']))
        story.append(Spacer(1, 14))

        # Section F — Additional information logged for completeness
        story.append(Paragraph("Section F — Additional information noted", section_heading_style))
        story.append(Spacer(1, 8))
        f_sec = _sec("section_f_excluded_from_scoring")
        if f_sec.get("section_intro"):
            story.append(Paragraph(f_sec.get("section_intro"), styles['Normal']))
            story.append(Spacer(1, 8))
        for item in f_sec.get("information_items", []) or []:
            story.append(Paragraph(f"• {item}", styles['Normal']))
        story.append(Spacer(1, 14))

        # Section G — Where the comparisons came from
        story.append(Paragraph("Section G — Where the comparisons came from", section_heading_style))
        story.append(Spacer(1, 8))
        g = _sec("section_g_traceability_appendix")
        if g.get("section_intro"):
            story.append(Paragraph(g.get("section_intro"), styles['Normal']))
            story.append(Spacer(1, 8))
        map_rows = g.get("mapping_table", []) or []
        table_data = [[_p("Finding / check"), _p("Source clause"), _p("Source type"), _p("Validation basis")]]
        for r in map_rows:
            r = _safe_row(r)
            table_data.append([
                _p(r.get("finding_or_check", ""), 60),
                _p(r.get("source_clause", ""), 50),
                _p(r.get("source_type", ""), 24),
                _p(r.get("validation_basis", ""), 30),
            ])
        if len(table_data) > 1:
            t = Table(table_data, colWidths=[1.5*inch, 1.5*inch, 1*inch, 2*inch], repeatRows=1)
            t.setStyle(_validation_table_style())
            story.append(t)
        story.append(Spacer(1, 8))
        methodology = (g.get('methodology_summary') or '')[:800]
        story.append(Paragraph(f"<i>Methodology:</i> {methodology}", styles['Normal']))
        story.append(Spacer(1, 14))

        # Section E — Schedule Float Profile (presentation only; does not affect acceptability)
        if float_stats:
            story.append(Paragraph("Section E — Schedule Float Profile", section_heading_style))
            story.append(Spacer(1, 8))
            story.append(Paragraph("<b>Summary Statistics</b>", styles['Heading3']))
            story.append(Spacer(1, 6))
            total = float_stats.get("total_activities", 0)
            critical_count = float_stats.get("critical_count", 0)
            critical_pct = float_stats.get("critical_pct", 0)
            under_40 = float_stats.get("under_40_count", 0)
            under_40_pct = float_stats.get("under_40_pct", 0)
            over_40 = float_stats.get("over_40_count", 0)
            over_40_pct = float_stats.get("over_40_pct", 0)
            max_f = float_stats.get("max_float", 0)
            med_f = float_stats.get("median_float", 0)
            story.append(Paragraph(
                f"• Total activities analysed: {total}<br/>"
                f"• Critical (0 float): {critical_count} ({critical_pct}%)<br/>"
                f"• Under 40 days: {under_40} ({under_40_pct}%)<br/>"
                f"• 40 days or more: {over_40} ({over_40_pct}%)<br/>"
                f"• Maximum float observed: {max_f} days<br/>"
                f"• Median float: {med_f} days",
                styles['Normal']
            ))
            story.append(Spacer(1, 12))
            story.append(Paragraph("<b>Float Distribution</b>", styles['Heading3']))
            story.append(Spacer(1, 6))
            hist_buf = _render_float_histogram_png(float_stats)
            if hist_buf and REPORTLAB_AVAILABLE and Image is not None:
                try:
                    hist_buf.seek(0)
                    # Use temp file for ReportLab Image (more reliable than BytesIO in some environments)
                    with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as _tmp:
                        _tmp.write(hist_buf.read())
                        histogram_temp_path = _tmp.name
                    img = Image(histogram_temp_path, width=4.75*inch, height=2.8*inch)
                    story.append(img)
                    story.append(Spacer(1, 10))
                except Exception as _e:
                    import logging
                    logging.getLogger(__name__).warning("[Schedule Float] Failed to embed histogram image: %s", _e)
                    if histogram_temp_path and os.path.exists(histogram_temp_path):
                        try:
                            os.unlink(histogram_temp_path)
                        except Exception:
                            pass
                    histogram_temp_path = None
            elif not hist_buf and float_stats:
                import logging
                logging.getLogger(__name__).info("[Schedule Float] Histogram not generated (install matplotlib to show graph).")
            story.append(Paragraph(
                "The float distribution illustrates how scheduling flexibility is allocated across the programme. "
                "A small proportion of low-float activities indicates concentrated delivery control, while high-float activities reflect sequencing flexibility.",
                styles['Normal']
            ))
            story.append(Paragraph(
                f"The programme is driven by {critical_count} critical activities, representing {critical_pct}% of the schedule.",
                styles['Normal']
            ))
            story.append(Spacer(1, 20))

        # Section H — Next steps
        story.append(Paragraph("Section H — Next steps", section_heading_style))
        story.append(Spacer(1, 8))
        h = _sec("section_h_next_steps")
        for action in h.get("next_steps", []) or []:
            story.append(Paragraph(f"• {action}", styles['Normal']))
        story.append(Spacer(1, 14))

        # What to review next (guidance only)
        what = report.get("section_what_to_review_next") or {}
        if what.get("items"):
            story.append(Paragraph("What to review next", section_heading_style))
            story.append(Spacer(1, 6))
            if what.get("section_intro"):
                story.append(Paragraph(what.get("section_intro"), styles['Normal']))
                story.append(Spacer(1, 6))
            for item in what.get("items", [])[:5]:
                story.append(Paragraph(f"• {item}", styles['Normal']))
            story.append(Spacer(1, 14))

        # User confirmations and notes (professional judgement; clearly separated)
        ucn = report.get("section_user_confirmations_and_notes") or {}
        ucn_items = ucn.get("items") or []
        if ucn_items:
            story.append(Paragraph("User confirmations and notes", section_heading_style))
            story.append(Spacer(1, 6))
            if ucn.get("section_intro"):
                story.append(Paragraph(ucn.get("section_intro"), styles['Normal']))
                story.append(Spacer(1, 6))
            for row in ucn_items[:30]:
                r = _safe_row(row)
                ts = (r.get("timestamp") or "")[:25]
                note = (r.get("note") or "")[:500]
                story.append(Paragraph(f"• [{ts}] {note}", styles['Normal']))
            story.append(Spacer(1, 14))

        # Appendix A — Detailed Float Analysis (critical + top 5 highest only; no full dump)
        appendix_float = report.get("appendix_float_trimmed") or []
        if appendix_float:
            story.append(Paragraph("Appendix A — Detailed Float Analysis", section_heading_style))
            story.append(Spacer(1, 8))
            story.append(Paragraph("Critical activities (0 float) and top 5 highest float activities.", styles['Normal']))
            story.append(Spacer(1, 6))
            table_data = [[_p("Activity Name"), _p("Total Float (days)"), _p("Critical (Yes/No)")]]
            for row in appendix_float:
                r = _safe_row(row)
                table_data.append([
                    _p(r.get("activity_name", ""), 100),
                    _p(r.get("total_float_days", ""), 22),
                    _p("Yes" if r.get("critical") else "No", 12),
                ])
            if len(table_data) > 1:
                t = Table(table_data, colWidths=[3.2*inch, 1.2*inch, 1.0*inch], repeatRows=1)
                t.setStyle(_validation_table_style())
                story.append(t)
            story.append(Spacer(1, 14))

        # Footer with version (branding only)
        footer_style = ParagraphStyle('FooterVersion', parent=styles['Normal'], fontSize=8, textColor=colors.grey, spaceBefore=24)
        story.append(Paragraph("NEC Engineering Analysis Report v1.0 — Programme validation output. Acceptability is determined under NEC Clause 31.", footer_style))

        doc.build(story)
        if histogram_temp_path and os.path.exists(histogram_temp_path):
            try:
                os.unlink(histogram_temp_path)
            except Exception:
                pass
        pdf_bytes = buffer.getvalue()
        buffer.close()
        if output_path:
            output_path = str(Path(output_path).resolve())
            Path(output_path).parent.mkdir(parents=True, exist_ok=True)
            with open(output_path, 'wb') as f:
                f.write(pdf_bytes)
            return output_path
        return pdf_bytes

    def _generate_validation_docx(
        self,
        json_data: Dict[str, Any],
        output_path: Optional[str] = None,
        report_options: Optional[Dict[str, Any]] = None,
    ) -> Union[bytes, str]:
        """Generate Programme Validation Report (Sections A–G) as DOCX."""
        from app.reporting.validation_report_builder import build_validation_report
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not available")
        report = build_validation_report(json_data)
        doc = Document()
        meta = _safe_dict(report.get("metadata"))
        doc.add_heading("Programme Validation Report", 0)
        doc.add_paragraph(f"Contract: {meta.get('contract_file', '—')}  |  Programme: {meta.get('xer_file', '—')}")
        doc.add_heading("Section A — Executive Summary", 1)
        a = _safe_dict(report.get("section_a_executive_summary"))
        doc.add_paragraph(f"Acceptability status: {a.get('acceptability_status', '—')}")
        doc.add_paragraph(f"Acceptability score: {a.get('acceptability_score', 0)}%")
        doc.add_paragraph(f"Quality / confidence score: {a.get('quality_score', 0)}%")
        doc.add_paragraph((a.get("verdict_paragraph") or "").replace("**", ""))
        if a.get("pm_note"):
            doc.add_paragraph(a.get("pm_note"), style="Intense Quote")
        doc.add_heading("Section B — What Determined the Outcome", 1)
        b = _safe_dict(report.get("section_b_what_determined_outcome"))
        if b.get("section_intro"):
            doc.add_paragraph(b.get("section_intro"))
        for tier_label, items in [("Tier 1 (Critical)", b.get("tier1_items", [])), ("Tier 2 (Quality / Risk)", b.get("tier2_items", [])), ("Tier 3 (Informational only)", b.get("tier3_items", []))]:
            if not items:
                continue
            doc.add_heading(tier_label, 2)
            table = doc.add_table(rows=1, cols=4)
            table.style = "Light Grid Accent 1"
            h = table.rows[0].cells
            h[0].text, h[1].text, h[2].text, h[3].text = "Requirement", "Outcome", "Source clause", "Explanation"
            for row in items[:20]:
                rw = _safe_row(row)
                r = table.add_row().cells
                r[0].text = str(rw.get("requirement", ""))[:50]
                r[1].text = str(rw.get("outcome", ""))[:20]
                r[2].text = str(rw.get("source_clause", ""))[:30]
                r[3].text = str(rw.get("explanation", ""))[:60]
        scope_sec = _safe_dict(report.get("section_scope_contract_alignment"))
        doc.add_heading("How scope and constraints are assessed", 1)
        doc.add_paragraph(
            "The programme is reviewed against the contract scope and constraints. "
            "Each contractual requirement is checked to confirm whether it is clearly represented in the programme activities or sequencing. "
            "Where a requirement is not represented, it is highlighted below."
        )
        doc.add_paragraph("Confidence level", style="Heading 3")
        doc.add_paragraph(
            "Confidence levels reflect how clearly programme activities demonstrate alignment with contractual obligations:"
        )
        doc.add_paragraph("High – Direct, explicit programme evidence supports the obligation.", style="List Bullet")
        doc.add_paragraph("Medium – Programme evidence is present but indirect or inferred.", style="List Bullet")
        doc.add_paragraph("Low – Limited or unclear programme evidence.", style="List Bullet")
        scope_rows_raw = scope_sec.get("scope_rows", []) or []
        scope_rows = []
        for row in scope_rows_raw:
            if isinstance(row, dict):
                scope_rows.append(row)
            else:
                scope_rows.append({
                    "contract_scope": str(row),
                    "programme_activities": [],
                    "representation_status": "Not specified",
                    "notes": "",
                })
        if scope_rows:
            table = doc.add_table(rows=1, cols=4)
            table.style = "Light Grid Accent 1"
            header = table.rows[0].cells
            header[0].text, header[1].text, header[2].text, header[3].text = (
                "Contract scope item",
                "Programme activities",
                "Coverage strength",
                "Notes",
            )
            for row in scope_rows[:25]:
                r = table.add_row().cells
                r[0].text = str(row.get("contract_scope", ""))[:120]
                pa = row.get("programme_activities", [])
                pa = pa if isinstance(pa, list) else ([pa] if pa is not None else [])
                programme_evidence = "\n".join(f"• {x}" for x in pa if x) or "—"
                r[1].text = programme_evidence[:500]
                r[2].text = str(row.get("representation_status", ""))[:50]
                r[3].text = _presentation_text(row.get("notes", ""))[:160]
        al_notes = scope_sec.get("activity_load_notes") or []
        if al_notes:
            doc.add_paragraph("Activity load transparency", style="Heading 3")
            for note in al_notes[:10]:
                doc.add_paragraph(note, style="List Bullet")
        obligations_evidenced_list = scope_sec.get("obligations_evidenced_list") or []
        obligations_explicit_assumption_list = scope_sec.get("obligations_explicit_assumption_list") or []
        obligations_not_represented_but_mandatory_list = scope_sec.get("obligations_not_represented_but_mandatory_list") or []
        if obligations_evidenced_list:
            doc.add_paragraph("Obligations evidenced by programme", style="Heading 3")
            doc.add_paragraph("The following obligations are evidenced by programme activities or constraints.")
            for ob in obligations_evidenced_list[:30]:
                o = ob if isinstance(ob, dict) else {}
                doc.add_paragraph(f"• [{o.get('id', '')}] {o.get('text', '')}", style="List Bullet")
        if obligations_explicit_assumption_list:
            doc.add_paragraph("Obligations with explicit assumption", style="Heading 3")
            doc.add_paragraph("Aligned via explicit assumption (client responsibility or out of scope at this stage).")
            for ob in obligations_explicit_assumption_list[:30]:
                o = ob if isinstance(ob, dict) else {}
                doc.add_paragraph(f"• [{o.get('id', '')}] {o.get('text', '')}", style="List Bullet")
                if o.get("exemption_reason"):
                    p = doc.add_paragraph()
                    p.add_run(f"  {o.get('exemption_reason')}").italic = True
        obligations_covered_by_later_submission_list = scope_sec.get("obligations_covered_by_later_submission_list") or []
        if obligations_covered_by_later_submission_list:
            doc.add_paragraph("Obligations assumed to be covered by later submission (non-blocking advisory only)", style="Heading 3")
            doc.add_paragraph("The following have only a 'covered by later submission' assumption. They do not set aligned and must never justify acceptance.")
            for ob in obligations_covered_by_later_submission_list[:30]:
                o = ob if isinstance(ob, dict) else {}
                doc.add_paragraph(f"• [{o.get('id', '')}] {o.get('text', '')}", style="List Bullet")
        if obligations_not_represented_but_mandatory_list:
            doc.add_paragraph("Not represented but mandatory", style="Heading 3")
            doc.add_paragraph("These mandatory obligations are not represented and block acceptability. They require either explicit programme activities or an explicit assumption.")
            for ob in obligations_not_represented_but_mandatory_list[:30]:
                o = ob if isinstance(ob, dict) else {}
                doc.add_paragraph(f"• [{o.get('id', '')}] {o.get('text', '')}", style="List Bullet")
        # Requires governance / future submission section removed from report output.
        if scope_sec.get("constraint_summary"):
            doc.add_paragraph(scope_sec.get("constraint_summary"))
        constraint_rows_raw = scope_sec.get("constraint_rows", []) or []
        constraint_rows = []
        for row in constraint_rows_raw:
            if isinstance(row, dict):
                constraint_rows.append(row)
            else:
                constraint_rows.append({
                    "constraint": str(row),
                    "programme_evidence": "",
                    "handling": "Implicitly managed",
                })
        if constraint_rows:
            table = doc.add_table(rows=1, cols=3)
            table.style = "Light Grid Accent 1"
            header = table.rows[0].cells
            header[0].text, header[1].text, header[2].text = (
                "Constraint",
                "Programme evidence",
                "Handling",
            )
            for row in constraint_rows[:25]:
                r = table.add_row().cells
                r[0].text = str(row.get("constraint", ""))[:120]
                r[1].text = str(row.get("programme_evidence", ""))[:160]
                r[2].text = str(row.get("handling", ""))[:40]
        if scope_sec.get("acceptability_clarification"):
            doc.add_paragraph(scope_sec.get("acceptability_clarification"))
        if scope_sec.get("reassurance"):
            doc.add_paragraph(scope_sec.get("reassurance"))

        doc.add_heading("Section C — Programme vs Contract Dates", 1)
        c = _safe_dict(report.get("section_c_dates"))
        if c.get("section_intro"):
            doc.add_paragraph(c.get("section_intro"))
        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"
        table.rows[0].cells[0].text, table.rows[0].cells[1].text = "Item", "Contract"
        table.rows[0].cells[2].text, table.rows[0].cells[3].text = "Programme", "Notes"
        for r in c.get("rows", []) or []:
            rw = _safe_row(r)
            row = table.add_row()
            row.cells[0].text = str(rw.get("item", ""))[:30]
            row.cells[1].text = str(rw.get("contract", ""))[:25]
            row.cells[2].text = str(rw.get("programme", ""))[:25]
            row.cells[3].text = _presentation_text(rw.get("notes", ""))[:40]
        doc.add_paragraph(f"Variance explanation: {_presentation_text(c.get('variance_explanation', ''))[:600]}")
        # Section D (Required activities) removed — not included in report output.
        doc.add_heading("Section E — Programme Quality & Realism (Advisory)", 1)
        e = _safe_dict(report.get("section_e_quality_realism_advisory"))
        if e.get("section_intro"):
            doc.add_paragraph(e.get("section_intro"))
        doc.add_paragraph(e.get("non_fatal_label", ""))
        for obs in e.get("sequencing_observations", [])[:10] or []:
            doc.add_paragraph(obs, style="List Bullet")
        doc.add_heading("Section F — Items Explicitly Excluded from Scoring", 1)
        f_sec = _safe_dict(report.get("section_f_excluded_from_scoring"))
        if f_sec.get("section_intro"):
            doc.add_paragraph(f_sec.get("section_intro"))
        for ex in f_sec.get("excluded_items", []) or []:
            exw = _safe_row(ex)
            doc.add_paragraph(f"• {exw.get('item', '')}: {exw.get('reason', '')}", style="List Bullet")
        doc.add_paragraph(f_sec.get("explicit_statement", ""))
        doc.add_heading("Section G — Traceability & Audit Appendix", 1)
        g = _safe_dict(report.get("section_g_traceability_appendix"))
        if g.get("section_intro"):
            doc.add_paragraph(g.get("section_intro"))
        doc.add_paragraph(g.get("methodology_summary", "")[:1000])
        if output_path:
            output_path = str(Path(output_path).resolve())
            Path(output_path).parent.mkdir(parents=True, exist_ok=True)
            doc.save(output_path)
            return output_path
        buffer = BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    def _generate_validation_html(
        self,
        json_data: Dict[str, Any],
        report_options: Optional[Dict[str, Any]] = None,
    ) -> str:
        """Generate Programme Validation Report (Sections A–G) as HTML."""
        from app.reporting.validation_report_builder import build_validation_report
        report = build_validation_report(json_data)
        meta = _safe_dict(report.get("metadata"))
        a = _safe_dict(report.get("section_a_executive_summary"))
        b = _safe_dict(report.get("section_b_what_determined_outcome"))
        c = _safe_dict(report.get("section_c_dates"))
        e = _safe_dict(report.get("section_e_quality_realism_advisory"))
        f_sec = _safe_dict(report.get("section_f_excluded_from_scoring"))
        g = _safe_dict(report.get("section_g_traceability_appendix"))
        html_parts = [
            "<!DOCTYPE html><html lang=\"en\"><head><meta charset=\"UTF-8\"><title>Programme Validation Report</title>",
            "<style>body{font-family:Segoe UI,Arial,sans-serif;margin:40px;color:#222;}",
            "h1{color:#003366;} h2{margin-top:28px;color:#003366;} h3{margin-top:16px;}",
            "table{width:100%;border-collapse:collapse;margin:12px 0;} th,td{border:1px solid #ccc;padding:8px;text-align:left;}",
            "th{background:#e6eef7;} .verdict{background:#f5f5f5;padding:12px;margin:12px 0;} .meta{color:#666;font-size:0.95em;}</style></head><body>",
            f"<h1>Programme Validation Report</h1>",
            f"<p class=\"meta\">Contract: {meta.get('contract_file', '—')} &nbsp;|&nbsp; Programme: {meta.get('xer_file', '—')}</p>",
            "<h2>Section A — Executive Summary</h2>",
            f"<p><strong>Acceptability status:</strong> {a.get('acceptability_status', '—')}</p>",
            f"<p><strong>Acceptability score:</strong> {a.get('acceptability_score', 0)}%</p>",
            f"<p><strong>Quality / confidence score:</strong> {a.get('quality_score', 0)}%</p>",
            f"<div class=\"verdict\"><p>{(a.get('verdict_paragraph') or '').replace('**', '')}</p></div>",
            f"<p class=\"pm-note\"><em>{_escape(a.get('pm_note') or '')}</em></p>" if a.get("pm_note") else "",
            "<h2>Section B — What Determined the Outcome</h2>",
            f"<p>{_escape(b.get('section_intro') or '')}</p>" if b.get("section_intro") else "",
        ]
        for tier_label, items in [("Tier 1 (Critical)", b.get("tier1_items", [])), ("Tier 2 (Quality / Risk)", b.get("tier2_items", [])), ("Tier 3 (Informational only)", b.get("tier3_items", []))]:
            if not items:
                continue
            html_parts.append(f"<h3>{tier_label}</h3><table><tr><th>Requirement</th><th>Outcome</th><th>Source clause</th><th>Explanation</th></tr>")
            for row in items:
                rw = _safe_row(row)
                html_parts.append(f"<tr><td>{_escape(rw.get('requirement',''))}</td><td>{_escape(rw.get('outcome',''))}</td><td>{_escape(rw.get('source_clause',''))}</td><td>{_escape(rw.get('explanation',''))[:80]}</td></tr>")
            html_parts.append("</table>")
        scope_sec = _safe_dict(report.get("section_scope_contract_alignment"))
        html_parts.append("<h2>How scope and constraints are assessed</h2>")
        html_parts.append(
            "<p>The programme is reviewed against the contract scope and constraints. "
            "Each contractual requirement is checked to confirm whether it is clearly represented in the programme activities or sequencing. "
            "Where a requirement is not represented, it is highlighted below.</p>"
        )
        html_parts.append("<h3>Confidence level</h3>")
        html_parts.append(
            "<p>Confidence levels reflect how clearly programme activities demonstrate alignment with contractual obligations:</p>"
            "<ul><li>High – Direct, explicit programme evidence supports the obligation.</li>"
            "<li>Medium – Programme evidence is present but indirect or inferred.</li>"
            "<li>Low – Limited or unclear programme evidence.</li></ul>"
        )
        scope_rows_raw = scope_sec.get("scope_rows", []) or []
        scope_rows = []
        for row in scope_rows_raw:
            if isinstance(row, dict):
                scope_rows.append(row)
            else:
                scope_rows.append({
                    "contract_scope": str(row),
                    "programme_activities": [],
                    "representation_status": "Not specified",
                    "notes": "",
                })
        if scope_rows:
            html_parts.append("<h3>Contract scope evidence</h3><table><tr><th>Contract scope item</th><th>Programme activities</th><th>Coverage strength</th><th>Notes</th></tr>")
            for row in scope_rows:
                pa = row.get("programme_activities", [])
                pa = pa if isinstance(pa, list) else ([pa] if pa is not None else [])
                programme_evidence = "<br/>".join(f"• {_escape(x)}" for x in pa if x) or "—"
                html_parts.append(
                    "<tr>"
                    f"<td>{_escape(row.get('contract_scope',''))}</td>"
                    f"<td>{programme_evidence}</td>"
                    f"<td>{_escape(row.get('representation_status',''))}</td>"
                    f"<td>{_escape(_presentation_text(row.get('notes','')))}</td>"
                    "</tr>"
                )
            html_parts.append("</table>")
            al_notes = scope_sec.get("activity_load_notes") or []
            if al_notes:
                html_parts.append("<h4>Activity load transparency</h4><ul>")
                for note in al_notes[:10]:
                    html_parts.append(f"<li>{_escape(note)}</li>")
                html_parts.append("</ul>")
        if scope_sec.get("constraint_summary"):
            html_parts.append(f"<p>{_escape(scope_sec.get('constraint_summary'))}</p>")
        constraint_rows_raw = scope_sec.get("constraint_rows", []) or []
        constraint_rows = []
        for row in constraint_rows_raw:
            if isinstance(row, dict):
                constraint_rows.append(row)
            else:
                constraint_rows.append({
                    "constraint": str(row),
                    "programme_evidence": "",
                    "handling": "Implicitly managed",
                })
        if constraint_rows:
            html_parts.append("<h3>Constraint coverage</h3><table><tr><th>Constraint</th><th>Programme evidence</th><th>Handling</th></tr>")
            for row in constraint_rows:
                html_parts.append(
                    "<tr>"
                    f"<td>{_escape(row.get('constraint',''))}</td>"
                    f"<td>{_escape(row.get('programme_evidence',''))}</td>"
                    f"<td>{_escape(row.get('handling',''))}</td>"
                    "</tr>"
                )
            html_parts.append("</table>")
        if scope_sec.get("acceptability_clarification"):
            html_parts.append(f"<p><strong>{_escape(scope_sec.get('acceptability_clarification'))}</strong></p>")
        if scope_sec.get("reassurance"):
            html_parts.append(f"<p>{_escape(scope_sec.get('reassurance'))}</p>")

        html_parts.append("<h2>Section C — Programme vs Contract Dates</h2>")
        if c.get("section_intro"):
            html_parts.append(f"<p>{_escape(c.get('section_intro'))}</p>")
        html_parts.append("<table><tr><th>Item</th><th>Contract</th><th>Programme</th><th>Notes</th></tr>")
        for r in c.get("rows", []) or []:
            rw = _safe_row(r)
            html_parts.append(f"<tr><td>{_escape(rw.get('item',''))}</td><td>{_escape(rw.get('contract',''))}</td><td>{_escape(rw.get('programme',''))}</td><td>{_escape(_presentation_text(rw.get('notes','')))[:50]}</td></tr>")
        html_parts.append("</table>")
        html_parts.append(f"<p><em>Variance explanation:</em> {_escape(_presentation_text(c.get('variance_explanation','')))}</p>")
        # Section D (Required activities) removed — not included in report output.
        html_parts.append("<h2>Section E — Programme Quality & Realism (Advisory)</h2>")
        if e.get("section_intro"):
            html_parts.append(f"<p>{_escape(e.get('section_intro'))}</p>")
        html_parts.append(f"<p><strong>NON-FATAL.</strong> {_escape(e.get('non_fatal_label',''))}</p><ul>")
        for obs in e.get("sequencing_observations", []):
            html_parts.append(f"<li>{_escape(obs)}</li>")
        for inf in e.get("interpretive_findings", []):
            html_parts.append(f"<li>{_escape(inf)}</li>")
        html_parts.append("</ul>")
        html_parts.append("<h2>Section F — Items Explicitly Excluded from Scoring</h2>")
        if f_sec.get("section_intro"):
            html_parts.append(f"<p>{_escape(f_sec.get('section_intro'))}</p>")
        html_parts.append("<ul>")
        for ex in f_sec.get("excluded_items", []) or []:
            exw = _safe_row(ex)
            html_parts.append(f"<li>{_escape(exw.get('item',''))}: {_escape(exw.get('reason',''))}</li>")
        html_parts.append(f"</ul><p><em>{_escape(f_sec.get('explicit_statement',''))}</em></p>")
        html_parts.append("<h2>Section G — Traceability & Audit Appendix</h2>")
        if g.get("section_intro"):
            html_parts.append(f"<p>{_escape(g.get('section_intro'))}</p>")
        html_parts.append("<table><tr><th>Finding / check</th><th>Source clause</th><th>Source type</th><th>Validation basis</th></tr>")
        for r in g.get("mapping_table", []) or []:
            rw = _safe_row(r)
            html_parts.append(f"<tr><td>{_escape(rw.get('finding_or_check',''))}</td><td>{_escape(rw.get('source_clause',''))}</td><td>{_escape(rw.get('source_type',''))}</td><td>{_escape(rw.get('validation_basis',''))}</td></tr>")
        html_parts.append("</table>")
        html_parts.append(f"<p><em>Methodology:</em> {_escape(g.get('methodology_summary',''))}</p>")
        html_parts.append("</body></html>")
        return "\n".join(html_parts)

    def _prepare_template_data(self, json_data: Dict[str, Any]) -> Dict[str, Any]:
        """
        Prepare JSON data for template rendering.
        Ensures all required fields exist with defaults.
        
        Handles both contract analysis JSON and validation JSON structures.
        """
        # Check if this is a validation JSON (has alignment and contract_summary)
        is_validation_json = "alignment" in json_data and "contract_summary" in json_data
        
        if is_validation_json:
            # This is a validation JSON - use the new structure
            contract_summary = json_data.get('contract_summary', {})
            programme_summary = json_data.get('programme_summary', {})
            alignment = json_data.get('alignment', {})
            risk_summary = json_data.get('risk_summary', {})
            validation_summary = json_data.get('validation_summary', {})
            
            # Extract contract data from contract_summary (new structure)
            contract_dates = {
                'starting_date': contract_summary.get('starting_date', {}).get('value', ''),
                'completion_date': contract_summary.get('completion_date', {}).get('value', ''),
                'possession_dates': ''  # Will be extracted from alignment
            }
            
            # Extract from contract_summary clauses if available
            clauses = contract_summary.get('clauses', {})
            if clauses:
                contract_dates['starting_date'] = contract_dates['starting_date'] or clauses.get('3.1', {}).get('value', '')
                contract_dates['completion_date'] = contract_dates['completion_date'] or clauses.get('3.3', {}).get('value', '')
            
            # Extract programme requirements
            programme_requirements = {
                'submit_first_programme_within': contract_summary.get('submit_first_programme_within', {}).get('value', ''),
                'revised_programme_interval': contract_summary.get('revised_programme_interval', {}).get('value', '')
            }
            
            # Extract payment terms
            payment_terms = {
                'retention_percentage': contract_summary.get('retention_percentage', {}).get('value', ''),
                'assessment_interval': clauses.get('5.2', {}).get('value', '') if clauses else '',
                'payment_period': clauses.get('5.3', {}).get('value', '') if clauses else '',
                'bond_amount': clauses.get('5.6', {}).get('value', '') if clauses else ''
            }
            
            # Extract defects
            defects = {
                'defects_date': clauses.get('4.1', {}).get('value', '') if clauses else '',
                'defect_correction_period': clauses.get('4.2', {}).get('value', '') if clauses else '',
                'landscaping_maintenance_period': clauses.get('4.3', {}).get('value', '') if clauses else ''
            }
            
            # Extract weather data
            weather_data_obj = contract_summary.get('weather_data', {})
            weather_data = {
                'recording_location': weather_data_obj.get('recording_location', ''),
                'measurement_data': weather_data_obj.get('measurement_data', ''),
                'historical_records_source': weather_data_obj.get('historical_records_source', '')
            }
            
            # Extract delay damages
            delay_damages = contract_summary.get('delay_damages', {}).get('value', '')
            
            # Calculate contract completeness
            total_fields = len([k for k in contract_summary.keys() if k != 'clauses'])
            present_fields = sum(1 for k, v in contract_summary.items() 
                                if k != 'clauses' and isinstance(v, dict) and v.get('status') == 'present')
            filled_percentage = (present_fields / total_fields * 100) if total_fields > 0 else 0.0
            
            # Convert clauses to extracted_clauses format for compatibility
            extracted_clauses = {}
            if clauses:
                for clause_num, clause_data in clauses.items():
                    extracted_clauses[clause_num] = {
                        'title': clause_data.get('title', ''),
                        'value': clause_data.get('value', ''),
                        'status': clause_data.get('status', 'missing')
                    }
            
            return {
                'project': json_data.get('metadata', {}).get('contract_file', 'Unknown Project'),
                'metadata': json_data.get('metadata', {}),
                'contract_completeness': {
                    'document_type': 'unknown',
                    'is_template': False,
                    'filled_percentage': filled_percentage,
                    'mandatory_missing': total_fields - present_fields,
                    'total_mandatory': total_fields,
                    'mandatory_filled': present_fields,
                    'mandatory_blank': 0,
                },
                'extracted_clauses': extracted_clauses,
                'contract_dates': contract_dates,
                'programme_requirements': programme_requirements,
                'delay_damages': delay_damages,
                'possession_dates': contract_dates.get('possession_dates', '—'),
                'defects': defects,
                'payment_terms': payment_terms,
                'weather_data': weather_data,
                'scope_items': [],
                'constraints': programme_summary.get('constraints', []),
                'milestones': programme_summary.get('list_of_milestones', []),
                # Add validation-specific data
                'validation_summary': validation_summary,
                'alignment': alignment,  # Use 'alignment' not 'nec_alignment'
                'risk_summary': risk_summary,  # Use 'risk_summary' not 'risks'
                'programme_summary': programme_summary,
                'contract_summary': contract_summary,  # Include full contract_summary
            }
        else:
            # This is a contract analysis JSON - use existing logic
            return {
                'project': json_data.get('project', 'Unknown Project'),
                'metadata': json_data.get('metadata', {}),
                'contract_completeness': json_data.get('contract_completeness', {
                    'document_type': 'unknown',
                    'is_template': False,
                    'filled_percentage': 0.0,
                    'mandatory_missing': 0,
                    'total_mandatory': 0,
                    'mandatory_filled': 0,
                    'mandatory_blank': 0,
                }),
                'extracted_clauses': json_data.get('extracted_clauses', {}),
                'contract_dates': json_data.get('contract_dates', {}),
                'programme_requirements': json_data.get('programme_requirements', {}),
                'delay_damages': json_data.get('delay_damages', '—'),
                'possession_dates': json_data.get('possession_dates', '—'),
                'defects': json_data.get('defects', {}),
                'payment_terms': json_data.get('payment_terms', {}),
                'weather_data': json_data.get('weather_data', {}),
                'scope_items': json_data.get('scope_items', []),
                'constraints': json_data.get('constraints', []),
                'milestones': json_data.get('milestones', []),
            }
    
    def generate_html(
        self,
        json_data: Dict[str, Any],
        narrative: Optional[Dict[str, Any]] = None,
        report_options: Optional[Dict[str, Any]] = None,
    ) -> str:
        """
        Generate HTML report from JSON data.
        If input is validation JSON (alignment + contract_summary), produces Programme Validation Report (Sections A–G).
        """
        if "alignment" in json_data and "contract_summary" in json_data:
            return self._generate_validation_html(json_data, report_options=report_options)
        
        from app.reporting.narrative_builder import NarrativeBuilder
        
        if not narrative:
            builder = NarrativeBuilder()
            narrative = builder.build_narrative(json_data)
        
        # Load HTML template if available
        if JINJA2_AVAILABLE:
            template_dir = Path(__file__).parent / "templates"
            if template_dir.exists():
                env = Environment(loader=FileSystemLoader(str(template_dir)))
                template = env.get_template("default_report_template.html")
                return template.render(
                    extraction_data=json_data,
                    narrative=narrative
                )
        
        # Fallback: Generate simple HTML
        return self._generate_simple_html(json_data, narrative)
    
    def _generate_simple_html(
        self,
        json_data: Dict[str, Any],
        narrative: Dict[str, Any]
    ) -> str:
        """Generate simple HTML report."""
        html = """<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>NEC Contract Analysis Report</title>
    <style>
        body { font-family: Arial, sans-serif; margin: 40px; }
        h1, h2, h3 { color: #003366; }
        table { width: 100%; border-collapse: collapse; margin: 20px 0; }
        th, td { border: 1px solid #999; padding: 8px; text-align: left; }
        th { background-color: #e6eef7; }
    </style>
</head>
<body>
"""
        
        # Add content sections
        exec_summary = narrative.get("executive_summary", {})
        html += f"<h1>NEC Contract Analysis Report</h1>"
        html += f"<h2>Executive Summary</h2>"
        html += f"<p>{exec_summary.get('summary', '')}</p>"
        
        # Add other sections...
        
        html += """
</body>
</html>
"""
        return html
