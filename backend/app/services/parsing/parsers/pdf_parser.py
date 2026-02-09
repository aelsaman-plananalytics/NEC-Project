"""
Hybrid PDF extraction pipeline for NEC Contract Data Parts 1 & 2.

Uses:
- PDFPlumber for normal text extraction
- Camelot/PDFPlumber for table extraction
- Tesseract OCR for scanned pages (fallback)
"""

import re
import os
from typing import Optional, List, Dict, Any, Tuple
from pathlib import Path

import pdfplumber
from docx import Document

# Optional imports for OCR and advanced table extraction
try:
    import pytesseract
    from pdf2image import convert_from_path
    from PIL import Image
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    print("[PDF_PARSER] OCR dependencies not available. Install pytesseract, pdf2image, and Pillow for OCR support.")

try:
    import camelot
    CAMELOT_AVAILABLE = True
except ImportError:
    CAMELOT_AVAILABLE = False
    print("[PDF_PARSER] Camelot not available. Using PDFPlumber for table extraction.")

# Optional import for PyMuPDF (fitz) for improved text extraction
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False
    fitz = None
    print("[PDF_PARSER] PyMuPDF not available. Install PyMuPDF for improved text extraction.")


class DocumentParser:
    """
    Hybrid document parser for extracting all content from NEC contract documents.
    
    Supports:
    - PDF text extraction with layout preservation
    - Table extraction (Camelot + PDFPlumber)
    - OCR fallback for scanned pages
    - Section classification
    - 4-layer data extraction
    """
    
    # Section classification patterns
    SECTION_PATTERNS = {
        "contract_data_part_1": [
            r"contract\s+data\s+part\s+one",
            r"contract\s+data\s+part\s+1",
            r"cd\s+part\s+one",
            r"cd\s+part\s+1"
        ],
        "contract_data_part_2": [
            r"contract\s+data\s+part\s+two",
            r"contract\s+data\s+part\s+2",
            r"cd\s+part\s+two",
            r"cd\s+part\s+2"
        ],
        "works_information": [
            r"works\s+information",
            r"wi\s+",
            r"description\s+of\s+works"
        ],
        "site_information": [
            r"site\s+information",
            r"si\s+"
        ],
        "drawing_register": [
            r"schedule\s+of\s+drawings",
            r"drawing\s+schedule",
            r"drawing\s+register"
        ],
        "key_dates": [
            r"key\s+dates",
            r"section\s+3",
            r"3\.\s*time"
        ],
        "z_clauses": [
            r"z\s+clause",
            r"additional\s+conditions"
        ],
        "payment": [
            r"payment",
            r"section\s+5",
            r"5\.\s*payment"
        ],
        "programme": [
            r"programme",
            r"section\s+3",
            r"3\.\s*time"
        ]
    }
    
    @staticmethod
    def extract_text(file_path: str) -> str:
        """
        Legacy method: Extract clean normalized text (backward compatibility).
        """
        structured = DocumentParser.extract_hybrid(file_path)
        return structured.get("contract_text", "")
    
    @staticmethod
    def extract_structured(file_path: str) -> Dict[str, Any]:
        """
        Legacy method: Extract structured data (backward compatibility).
        Returns pages with text_blocks and tables.
        """
        hybrid = DocumentParser.extract_hybrid(file_path)
        
        # Convert to legacy format
        pages = []
        for page_data in hybrid.get("pages", []):
            pages.append({
                "page_no": page_data.get("page_number", 0),
                "text_blocks": page_data.get("text_blocks", []),
                "tables": page_data.get("tables", [])
            })
        
        return {"pages": pages}
    
    @staticmethod
    def extract_hybrid(file_path: str) -> Dict[str, Any]:
        """
        Hybrid PDF extraction with OCR fallback and table capture.
        
        Produces 4 layers of data:
        - Layer 1: Full cleaned text
        - Layer 2: Classified sections
        - Layer 3: Table data
        - Layer 4: Scope items (extracted separately)
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            dict: {
                "contract_text": "...",
                "sections": {...},
                "tables": [...],
                "key_dates": {...},
                "constraints": {...},
                "scope_items": [...],
                "drawings": [...],
                "metadata": {...},
                "pages": [...]
            }
        """
        file_path_obj = Path(file_path)
        
        if not file_path_obj.exists():
            raise FileNotFoundError(f"PDF file not found: {file_path}")
        
        if file_path_obj.suffix.lower() != '.pdf':
            raise ValueError(f"File must be a PDF: {file_path}")
        
        print(f"[PDF_PARSER] Starting hybrid extraction: {file_path}")
        
        # Initialize result structure
        result = {
            "contract_text": "",
            "sections": {},
            "tables": [],
            "key_dates": {},
            "constraints": {},
            "scope_items": [],
            "drawings": [],
            "metadata": {
                "file_name": file_path_obj.name,
                "total_pages": 0,
                "ocr_used": False,
                "tables_found": 0
            },
            "pages": []
        }
        
        try:
            with pdfplumber.open(file_path) as pdf:
                total_pages = len(pdf.pages)
                result["metadata"]["total_pages"] = total_pages
                
                all_text_blocks = []
                all_tables = []
                section_texts = {}
                
                for page_num, page in enumerate(pdf.pages, start=1):
                    print(f"[PDF_PARSER] Processing page {page_num}/{total_pages}")
                    
                    page_data = {
                        "page_number": page_num,
                        "text_blocks": [],
                        "tables": [],
                        "section": None,
                        "ocr_used": False
                    }
                    
                    # Extract text blocks (preserving layout)
                    words = page.extract_words(
                        x_tolerance=3,
                        y_tolerance=3,
                        keep_blank_chars=False
                    )
                    
                    text_content = ""
                    if words:
                        text_blocks = DocumentParser._group_words_into_blocks(words)
                        cleaned_blocks = [
                            DocumentParser._clean_text_block(block)
                            for block in text_blocks
                            if block.strip()
                        ]
                        page_data["text_blocks"] = cleaned_blocks
                        text_content = " ".join(cleaned_blocks)
                        all_text_blocks.extend(cleaned_blocks)
                    
                    # Check if OCR is needed (< 50 characters suggests scanned page)
                    if len(text_content) < 50 and OCR_AVAILABLE:
                        print(f"[PDF_PARSER] Low text content on page {page_num}, attempting OCR...")
                        ocr_text = DocumentParser._extract_with_ocr(file_path, page_num)
                        if ocr_text and len(ocr_text) > len(text_content):
                            text_content = ocr_text
                            page_data["text_blocks"] = [ocr_text]
                            page_data["ocr_used"] = True
                            result["metadata"]["ocr_used"] = True
                            all_text_blocks.append(ocr_text)
                    
                    # Extract tables
                    page_tables = DocumentParser._extract_tables_from_page(
                        file_path, page_num, page
                    )
                    # Normalize tables to simple list of lists format
                    # Each table is a list of rows, each row is a list of cells
                    normalized_tables = []
                    for table_dict in page_tables:
                        rows = table_dict.get("rows", [])
                        if rows:
                            simple_rows = []
                            for row in rows:
                                if isinstance(row, dict):
                                    # Convert dict row to list: [field, value] or [all values]
                                    if "field" in row and "value" in row:
                                        simple_rows.append([row["field"], row["value"]])
                                    else:
                                        simple_rows.append(list(row.values()))
                                elif isinstance(row, list):
                                    simple_rows.append(row)
                            if simple_rows:
                                normalized_tables.append(simple_rows)
                    
                    # Assign normalized tables to page (list of lists format)
                    page_data["tables"] = normalized_tables
                    all_tables.extend(page_tables)
                    
                    # Classify section
                    section_type = DocumentParser._classify_section(text_content)
                    page_data["section"] = section_type
                    if section_type:
                        if section_type not in section_texts:
                            section_texts[section_type] = []
                        section_texts[section_type].append(text_content)
                    
                    result["pages"].append(page_data)
                
                # Build Layer 1: Full cleaned text
                result["contract_text"] = "\n\n".join(all_text_blocks)
                
                # Build Layer 2: Classified sections
                result["sections"] = {
                    section: "\n\n".join(texts)
                    for section, texts in section_texts.items()
                }
                
                # Build Layer 3: Table data
                result["tables"] = all_tables
                result["metadata"]["tables_found"] = len(all_tables)
                
                # Extract key dates and constraints from tables
                result["key_dates"] = DocumentParser._extract_key_dates_from_tables(all_tables)
                result["constraints"] = DocumentParser._extract_constraints_from_tables(all_tables)
                
                # Extract drawings from tables
                result["drawings"] = DocumentParser._extract_drawings_from_tables(all_tables)
                
                print(f"[PDF_PARSER] Extraction complete:")
                print(f"  - Pages: {total_pages}")
                print(f"  - Tables: {len(all_tables)}")
                print(f"  - Sections: {len(result['sections'])}")
                print(f"  - OCR used: {result['metadata']['ocr_used']}")
                
                return result
                
        except Exception as e:
            raise Exception(f"Error extracting PDF: {e}")
    
    @staticmethod
    def _extract_tables_from_page(
        file_path: str,
        page_num: int,
        page: Any
    ) -> List[Dict[str, Any]]:
        """
        Extract tables from a page using Camelot (preferred) or PDFPlumber (fallback).
        
        Returns list of table dictionaries with structure:
        {
            "table_name": "...",
            "page_number": int,
            "rows": [{"field": "...", "value": "..."}, ...]
        }
        """
        tables = []
        
        # Try Camelot first (better for vector PDFs) - STREAM MODE ONLY
        if CAMELOT_AVAILABLE:
            try:
                camelot_tables = camelot.read_pdf(
                    str(file_path),
                    pages=str(page_num),
                    flavor="stream",              # <-- CRITICAL FIX: Use stream mode only
                    row_tol=10,                   # better row grouping
                    column_tol=20,                # detect column separation
                    split_text=True,
                    strip_text="\n",
                    edge_tol=500,                 # allow wide spacing
                )
                
                # Normalize Camelot table output into simple rows
                extracted_tables = []
                for t in camelot_tables:
                    df = t.df
                    rows = df.values.tolist()
                    cleaned = [[str(cell).strip() for cell in row] for row in rows]
                    extracted_tables.append(cleaned)
                
                # Convert normalized tables to our format
                # NEW: Return tables as list of lists format for table-first extraction
                for cleaned_table in extracted_tables:
                    if not cleaned_table or len(cleaned_table) == 0:
                        continue
                    
                    # Return table as list of lists: [["col1", "col2"], ["col1", "col2"], ...]
                    # This is the format expected by ContractDataTableExtractor
                    tables.append({
                        "table_name": DocumentParser._infer_table_name_from_data(cleaned_table, page_num),
                        "page_number": page_num,
                        "rows": cleaned_table  # Direct list-of-lists format
                    })
                
                if tables:
                    return tables
            except Exception as e:
                print(f"[PDF_PARSER] Camelot stream extraction failed for page {page_num}: {e}")
        
        # Fallback to PDFPlumber
        try:
            pdfplumber_tables = page.extract_tables()
            for table in pdfplumber_tables:
                if table and len(table) > 0:
                    # Return table as list of lists format
                    # Clean each cell and ensure it's a string
                    cleaned_table = []
                    for row in table:
                        cleaned_row = [str(cell).strip() if cell else "" for cell in row]
                        if cleaned_row:  # Only add non-empty rows
                            cleaned_table.append(cleaned_row)
                    
                    if cleaned_table:
                        tables.append({
                            "table_name": DocumentParser._infer_table_name_from_data(cleaned_table, page_num),
                            "page_number": page_num,
                            "rows": cleaned_table  # Direct list-of-lists format
                        })
        except Exception as e:
            print(f"[PDF_PARSER] PDFPlumber table extraction failed for page {page_num}: {e}")
        
        return tables
    
    @staticmethod
    def _infer_table_name(table_data: Any, page_num: int) -> str:
        """Infer table name from table data or page context."""
        # Try to get first row as header
        if hasattr(table_data, 'iloc'):  # Pandas DataFrame
            if len(table_data) > 0:
                first_row = " ".join(str(cell) for cell in table_data.iloc[0] if cell)
                if first_row:
                    # Check for common table names
                    first_row_lower = first_row.lower()
                    if "key date" in first_row_lower:
                        return "Key Dates"
                    if "starting date" in first_row_lower or "completion date" in first_row_lower:
                        return "Time Section"
                    if "drawing" in first_row_lower or "schedule" in first_row_lower:
                        return "Drawing Schedule"
                    if "insurance" in first_row_lower:
                        return "Insurance"
                    if "price" in first_row_lower or "cost" in first_row_lower:
                        return "Price List"
        elif isinstance(table_data, list) and len(table_data) > 0:
            first_row = " ".join(str(cell) for cell in table_data[0] if cell)
            if first_row:
                first_row_lower = first_row.lower()
                if "key date" in first_row_lower:
                    return "Key Dates"
                if "drawing" in first_row_lower:
                    return "Drawing Schedule"
        
        return f"Table on page {page_num}"
    
    @staticmethod
    def _infer_table_name_from_data(table: List[List], page_num: int) -> str:
        """Infer table name from PDFPlumber table data."""
        if not table or len(table) == 0:
            return f"Table on page {page_num}"
        
        first_row = " ".join(str(cell) for cell in table[0] if cell)
        first_row_lower = first_row.lower()
        
        if "key date" in first_row_lower:
            return "Key Dates"
        if "drawing" in first_row_lower or "schedule" in first_row_lower:
            return "Drawing Schedule"
        if "insurance" in first_row_lower:
            return "Insurance"
        if "starting date" in first_row_lower or "completion date" in first_row_lower:
            return "Time Section"
        
        return f"Table on page {page_num}"
    
    @staticmethod
    def _convert_table_to_rows(table_df: Any, table_name: str) -> List[Dict[str, Any]]:
        """Convert Camelot table DataFrame to row format."""
        rows = []
        
        if not hasattr(table_df, 'iloc'):
            return rows
        
        # Get headers from first row
        headers = []
        if len(table_df) > 0:
            headers = [str(cell).strip() for cell in table_df.iloc[0] if cell]
        
        # Convert data rows
        start_idx = 1 if headers else 0
        for idx in range(start_idx, len(table_df)):
            row_data = table_df.iloc[idx]
            if len(headers) == 2:
                # Two-column table: field-value pairs
                rows.append({
                    "field": str(row_data.iloc[0]).strip() if len(row_data) > 0 else "",
                    "value": str(row_data.iloc[1]).strip() if len(row_data) > 1 else ""
                })
            else:
                # Multi-column: combine all cells
                row_dict = {}
                for i, header in enumerate(headers):
                    if i < len(row_data):
                        row_dict[header] = str(row_data.iloc[i]).strip()
                if row_dict:
                    rows.append(row_dict)
        
        return rows
    
    @staticmethod
    def _convert_table_to_rows_pdfplumber(
        table: List[List],
        table_name: str
    ) -> List[Dict[str, Any]]:
        """Convert PDFPlumber table to row format."""
        rows = []
        
        if not table or len(table) == 0:
            return rows
        
        # Get headers from first row
        headers = [str(cell).strip() for cell in table[0] if cell]
        
        # Convert data rows
        start_idx = 1 if headers else 0
        for row in table[start_idx:]:
            if not row:
                continue
            
            if len(headers) == 2:
                # Two-column table: field-value pairs
                rows.append({
                    "field": str(row[0]).strip() if len(row) > 0 and row[0] else "",
                    "value": str(row[1]).strip() if len(row) > 1 and row[1] else ""
                })
            else:
                # Multi-column: combine all cells
                row_dict = {}
                for i, header in enumerate(headers):
                    if i < len(row) and row[i]:
                        row_dict[header] = str(row[i]).strip()
                if row_dict:
                    rows.append(row_dict)
        
        return rows
    
    @staticmethod
    def _extract_with_ocr(file_path: str, page_num: int) -> str:
        """Extract text from a page using OCR (Tesseract)."""
        if not OCR_AVAILABLE:
            return ""
        
        try:
            # Convert PDF page to image
            images = convert_from_path(file_path, first_page=page_num, last_page=page_num)
            if not images:
                return ""
            
            # Run OCR on the image
            ocr_text = pytesseract.image_to_string(images[0])
            return ocr_text.strip()
        except Exception as e:
            print(f"[PDF_PARSER] OCR failed for page {page_num}: {e}")
            return ""
    
    @staticmethod
    def _classify_section(text: str) -> Optional[str]:
        """Classify text into section types."""
        text_lower = text.lower()
        
        for section_type, patterns in DocumentParser.SECTION_PATTERNS.items():
            for pattern in patterns:
                if re.search(pattern, text_lower):
                    return section_type
        
        return None
    
    @staticmethod
    def _extract_key_dates_from_tables(tables: List[Dict[str, Any]]) -> Dict[str, Any]:
        """Extract key dates from tables."""
        key_dates = {
            "starting_date": "",
            "completion_date": "",
            "access_dates": [],
            "key_dates": []
        }
        
        for table in tables:
            if "key date" in table.get("table_name", "").lower():
                for row in table.get("rows", []):
                    field = row.get("field", "").lower()
                    value = row.get("value", "").strip()
                    
                    if "starting" in field:
                        key_dates["starting_date"] = value
                    elif "completion" in field:
                        key_dates["completion_date"] = value
                    elif "possession" in field or "access" in field:
                        key_dates["access_dates"].append(value)
                    elif "key date" in field:
                        key_dates["key_dates"].append(value)
        
        return key_dates
    
    @staticmethod
    def _extract_constraints_from_tables(tables: List[Dict[str, Any]]) -> Dict[str, Any]:
        """Extract constraints from tables."""
        constraints = {}
        
        for table in tables:
            table_name = table.get("table_name", "").lower()
            if "constraint" in table_name or "access" in table_name:
                for row in table.get("rows", []):
                    field = row.get("field", "").strip()
                    value = row.get("value", "").strip()
                    if field and value:
                        constraints[field] = value
        
        return constraints
    
    @staticmethod
    def _extract_drawings_from_tables(tables: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """Extract drawings from drawing schedule tables."""
        drawings = []
        
        for table in tables:
            table_name = table.get("table_name", "").lower()
            if "drawing" in table_name or "schedule" in table_name:
                for row in table.get("rows", []):
                    # Try to extract drawing code, title, discipline
                    drawing = {}
                    
                    if isinstance(row, dict):
                        if "code" in row or "number" in row:
                            drawing["code"] = row.get("code") or row.get("number", "")
                        if "title" in row:
                            drawing["title"] = row["title"]
                        if "discipline" in row:
                            drawing["discipline"] = row["discipline"]
                    elif isinstance(row, list) and len(row) >= 2:
                        drawing["code"] = str(row[0]).strip() if row[0] else ""
                        drawing["title"] = str(row[1]).strip() if len(row) > 1 and row[1] else ""
                    
                    if drawing.get("code") or drawing.get("title"):
                        drawings.append(drawing)
        
        return drawings
    
    @staticmethod
    def _clean_text(text: str) -> str:
        """Clean and normalize extracted text."""
        if not text:
            return ""
        
        # Remove CID artifacts
        text = re.sub(r'\(cid:\d+\)', '', text)
        
        # Remove null bytes and BOM
        text = re.sub(r'\x00', '', text)
        text = re.sub(r'\ufeff', '', text)
        
        # Normalize whitespace
        text = re.sub(r' +', ' ', text)
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        return text.strip()
    
    @staticmethod
    def _clean_text_block(text: str) -> str:
        """Clean a single text block."""
        if not text:
            return ""
        
        # Remove CID artifacts
        text = re.sub(r'\(cid:\d+\)', '', text)
        
        # Normalize whitespace
        text = re.sub(r'\s+', ' ', text)
        
        return text.strip()
    
    @staticmethod
    def _group_words_into_blocks(words: List[Dict[str, Any]], 
                                 line_tolerance: float = 5.0) -> List[str]:
        """Group extracted words into text blocks based on proximity."""
        if not words:
            return []
        
        sorted_words = sorted(words, key=lambda w: (w.get('top', 0), w.get('left', 0)))
        
        blocks = []
        current_line = []
        current_y = None
        
        for word in sorted_words:
            word_text = word.get('text', '').strip()
            if not word_text:
                continue
            
            word_y = word.get('top', 0)
            
            if current_y is None or abs(word_y - current_y) <= line_tolerance:
                current_line.append(word_text)
                if current_y is None:
                    current_y = word_y
            else:
                if current_line:
                    blocks.append(' '.join(current_line))
                current_line = [word_text]
                current_y = word_y
        
        if current_line:
            blocks.append(' '.join(current_line))
        
        return blocks
    
    @staticmethod
    def _extract_from_word(file_path: str) -> str:
        """Extract text from Word document (legacy method)."""
        try:
            doc = Document(file_path)
            if len(doc.paragraphs) == 0:
                raise ValueError("Word document contains no paragraphs")
            
            paragraph_texts = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraph_texts.append(text)
            
            for table in doc.tables:
                for row in table.rows:
                    row_texts = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_texts.append(cell_text)
                    if row_texts:
                        paragraph_texts.append(" | ".join(row_texts))
            
            if not paragraph_texts:
                raise ValueError("No text could be extracted from Word document")
            
            raw_text = "\n".join(paragraph_texts)
            return DocumentParser._clean_text(raw_text)
            
        except Exception as e:
            raise Exception(f"Error reading Word document: {e}")
    
    @staticmethod
    def extract_clean_text(pdf_path: str) -> str:
        """
        Extract clean, continuous text from PDF using PyMuPDF.
        
        Features:
        - Table detection via PyMuPDF
        - Multi-column layout handling
        - Removal of headers/footers
        - Correct joining of wrapped lines
        - Preservation of section numbers (1., 1.1, 1.2.3)
        - Output a single continuous clean text block
        
        Args:
            pdf_path: Path to PDF file
            
        Returns:
            str: Clean, continuous text with preserved section numbering
        """
        if not PYMUPDF_AVAILABLE:
            # Fallback to pdfplumber if PyMuPDF not available
            print("[PDF_PARSER] PyMuPDF not available, falling back to pdfplumber")
            return DocumentParser._extract_from_pdfplumber(pdf_path)
        
        try:
            doc = fitz.open(pdf_path)
            all_text = []
            
            for page_num, page in enumerate(doc, start=1):
                blocks = page.get_text("blocks")
                page_height = page.rect.height
                page_width = page.rect.width
                
                # Separate blocks by column (for two-column layouts)
                left_column = []
                right_column = []
                full_width = []
                
                for block in blocks:
                    if len(block) < 6:
                        continue
                    
                    x0, y0, x1, y1, text, block_no = block[:6]
                    
                    # Skip headers/footers (very small text at top/bottom)
                    # Headers typically at y0 < 50, footers at y1 > page_height - 50
                    if y0 < 50 or y1 > (page_height - 50):
                        continue
                    
                    # Skip empty lines
                    if not text or not text.strip():
                        continue
                    
                    # Detect column layout
                    block_center_x = (x0 + x1) / 2
                    mid_page = page_width / 2
                    
                    # PRESERVE line breaks around NEC labels - don't join everything
                    # Only normalize excessive whitespace, keep newlines for label detection
                    cleaned = text.strip()
                    # Normalize multiple spaces but preserve single newlines
                    cleaned = re.sub(r'[ \t]+', ' ', cleaned)  # Multiple spaces/tabs to single space
                    # Preserve newlines - they're important for label detection
                    
                    if cleaned:
                        # Classify by column position
                        if block_center_x < mid_page - 50:  # Left column
                            left_column.append((y0, x0, cleaned))
                        elif block_center_x > mid_page + 50:  # Right column
                            right_column.append((y0, x0, cleaned))
                        else:  # Full width (tables, headings)
                            full_width.append((y0, x0, cleaned))
                
                # Sort each column by vertical position (top to bottom)
                left_column.sort(key=lambda x: (x[0], x[1]))  # Sort by y, then x
                right_column.sort(key=lambda x: (x[0], x[1]))
                full_width.sort(key=lambda x: (x[0], x[1]))
                
                # Merge columns in reading order (left column first, then right column)
                # Use a simple merge: if two-column layout detected, read left then right
                page_lines = []
                
                if len(left_column) > 0 and len(right_column) > 0:
                    # Two-column layout: merge by reading order
                    # Read left column top-to-bottom, then right column top-to-bottom
                    page_lines.extend([line[2] for line in left_column])
                    page_lines.extend([line[2] for line in right_column])
                else:
                    # Single column or full-width: sort by y-coordinate only
                    all_blocks = left_column + right_column + full_width
                    all_blocks.sort(key=lambda x: x[0])
                    page_lines = [line[2] for line in all_blocks]
                
                # Also add full-width blocks (they may contain important table data)
                # But avoid duplicates
                for line in full_width:
                    if line[2] not in page_lines:
                        # Insert at appropriate position based on y-coordinate
                        inserted = False
                        for i, existing_line in enumerate(page_lines):
                            # Find approximate position (this is heuristic)
                            if not inserted:
                                page_lines.append(line[2])
                                inserted = True
                                break
                
                all_text.extend(page_lines)
            
            doc.close()
            
            # Join all pages with newlines preserved for label detection
            raw_text = "\n".join(all_text)
            
            # Apply PDF-level cleaning (removes TOC leaders, page numbers, fixes hyphenation)
            # BUT preserves line breaks around NEC labels
            final_text = DocumentParser.clean_raw_pdf_text(raw_text)
            
            print(f"[PDF_PARSER] Extracted {len(final_text)} characters of clean text")
            return final_text.strip()
            
        except Exception as e:
            print(f"[PDF_PARSER] PyMuPDF extraction failed: {e}, falling back to pdfplumber")
            # Fallback to pdfplumber
            return DocumentParser._extract_from_pdfplumber(pdf_path)
    
    @staticmethod
    def clean_raw_pdf_text(text: str) -> str:
        """
        Low-level mechanical cleaning of raw PDF text.
        
        Removes:
        - Dotted TOC leaders (e.g., "....................... 4")
        - Page numbers on isolated lines
        - Repeated headers/footers (already handled in extract_clean_text)
        - Excessive whitespace
        
        Fixes:
        - Hyphenation (e.g., "construc-\ntion" → "construction")
        - Wrapped lines that belong to the same sentence
        
        Preserves:
        - All letters, numbers, dates, clause references, table values
        - All symbols in clause content (%, /, :, (, ), etc.)
        
        Args:
            text: Raw text extracted from PDF
            
        Returns:
            str: Cleaned text with structural noise removed
        """
        if not text:
            return ""
        
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            original_line = line
            line = line.strip()
            
            # Skip empty lines
            if not line:
                continue
            
            # Remove dotted TOC leaders: "....................... 4" or ".... ............. .... 12"
            # Pattern: 5+ dots followed by optional spaces and a number at end
            line = re.sub(r'\.{5,}\s*\d+\b', '', line)
            
            # Remove page numbers on isolated lines
            # A line containing only a number or "Page X" should be removed
            if re.match(r'^(Page\s+)?\d+\s*$', line, re.IGNORECASE):
                continue
            
            # Normalize spacing (multiple spaces to single space)
            line = re.sub(r' +', ' ', line)
            
            # Skip lines that are only dots, slashes, underscores, or spaces
            if re.match(r'^[\s\.\/_\-]+$', line):
                continue
            
            cleaned_lines.append(line)
        
        # Second pass: merge hyphenated lines and wrapped lines intelligently
        merged_lines = []
        i = 0
        while i < len(cleaned_lines):
            current_line = cleaned_lines[i]
            
            # Fix hyphenation: if line ends with hyphen, merge with next line
            if current_line.endswith('-') and len(current_line) > 1 and i + 1 < len(cleaned_lines):
                # Remove trailing hyphen and merge with next line (no space)
                current_line = current_line.rstrip('-').rstrip()
                next_line = cleaned_lines[i + 1]
                current_line = current_line + next_line
                i += 2  # Skip next line since we merged it
                merged_lines.append(current_line)
                continue
            
            # Check if we should merge with next line (wrapped lines)
            # IMPROVED: Preserve line breaks around NEC labels
            if i + 1 < len(cleaned_lines):
                next_line = cleaned_lines[i + 1]
                
                # NEC label patterns - preserve line breaks around these
                nec_label_patterns = [
                    r'Starting\s+Date',
                    r'Completion\s+Date',
                    r'Access\s+Date[s]?',
                    r'Possession\s+Date[s]?',
                    r'Defects\s+Date',
                    r'Defect\s+Correction\s+Period',
                    r'Programme\s+submission',
                    r'revised\s+programme[s]?',
                    r'Delay\s+damages',
                    r'Retention',
                    r'Bond\s+Amount',
                    r'Assessment\s+Interval',
                    r'Payment\s+Period',
                    r'Key\s+Date[s]?',
                    r'Weather\s+Recording',
                    r'Weather\s+measurements?',
                ]
                
                # Check if current line contains an NEC label
                has_nec_label = any(re.search(pattern, current_line, re.IGNORECASE) for pattern in nec_label_patterns)
                
                # Merge if:
                # 1. Current line doesn't end with sentence-ending punctuation (. ! ?)
                # 2. Next line doesn't start with capital letter (unless it's a clause number)
                # 3. Current line doesn't end with colon (likely a heading)
                # 4. Current line doesn't contain an NEC label (preserve line breaks around labels)
                should_merge = (
                    not re.search(r'[.!?]\s*$', current_line) and
                    not current_line.endswith(':') and
                    not has_nec_label and  # NEW: Don't merge if line contains NEC label
                    not re.match(r'^\d+\.\d+', next_line) and  # Don't merge with clause numbers
                    not re.match(r'^[A-Z][a-z]+', next_line)  # Don't merge if next starts with capital
                )
                
                if should_merge:
                    # Merge lines with a space
                    current_line = current_line + ' ' + next_line
                    i += 2  # Skip next line since we merged it
                    merged_lines.append(current_line)
                    continue
            
            merged_lines.append(current_line)
            i += 1
        
        # Join all lines
        final_text = '\n'.join(merged_lines)
        
        # Final normalization: collapse multiple spaces
        final_text = re.sub(r' +', ' ', final_text)
        
        # Normalize newlines (max 2 consecutive)
        final_text = re.sub(r'\n{3,}', '\n\n', final_text)
        
        return final_text.strip()
    
    @staticmethod
    def _extract_from_pdfplumber(pdf_path: str) -> str:
        """
        Fallback extraction using pdfplumber.
        
        Args:
            pdf_path: Path to PDF file
            
        Returns:
            str: Extracted text
        """
        try:
            text_blocks = []
            with pdfplumber.open(pdf_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_blocks.append(page_text)
            raw_text = "\n".join(text_blocks)
            # Apply cleaning to pdfplumber output too
            return DocumentParser.clean_raw_pdf_text(raw_text)
        except Exception as e:
            raise Exception(f"Error extracting text from PDF: {e}")